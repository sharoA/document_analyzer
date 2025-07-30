#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
需求文档向量化初始化脚本
专门用于处理 /Users/renyu/Documents/knowledge_base/链数_LS/需求文档 目录下的DOCX文件
按标题分段进行向量化存储
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

# 文件处理相关
import docx
from sentence_transformers import SentenceTransformer

# Weaviate 和 LangChain
import weaviate
import weaviate.classes as wvc
from langchain.schema import Document

# 项目配置
try:
    from src.utils.weaviate_helper import get_weaviate_client
except ImportError:
    import sys
    project_root = os.path.dirname(os.path.abspath(__file__))
    sys.path.insert(0, project_root)
    from src.utils.weaviate_helper import get_weaviate_client

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class RequirementsDocsInitializer:
    """需求文档初始化器 - 简化版"""
    
    def __init__(self, requirements_docs_path: str = "/Users/renyu/Documents/knowledge_base/链数_LS/需求文档"):
        """
        初始化需求文档处理器
        
        Args:
            requirements_docs_path: 需求文档目录路径
        """
        self.requirements_docs_path = Path(requirements_docs_path)
        self.weaviate_client = None
        self.embeddings_model = None
        
        # 只处理DOCX文件
        self.supported_extensions = {'.docx'}
        
        # 初始化组件
        self._initialize_components()
    
    def _initialize_components(self):
        """初始化各种组件"""
        try:
            # 初始化 Weaviate 客户端
            self.weaviate_client = get_weaviate_client()
            logger.info("✅ Weaviate 客户端初始化成功")
            
            # 初始化嵌入模型 - 使用bge-large-zh（1024维，中文优化）
            self.embeddings_model = SentenceTransformer('BAAI/bge-large-zh')
            logger.info("✅ 嵌入模型加载成功 (bge-large-zh, 1024维，中文优化)")
            
            logger.info("✅ 组件初始化完成")
            
        except Exception as e:
            logger.error(f"❌ 组件初始化失败: {e}")
            raise
    
    def _clear_weaviate_database(self):
        """清空Weaviate数据库中的需求文档数据"""
        try:
            logger.info("🗑️ 开始清空需求文档相关数据...")
            
            collection_name = "Document"
            if self.weaviate_client.collections.exists(collection_name):
                # 删除项目为LS的文档
                collection = self.weaviate_client.collections.get(collection_name)
                
                # 查询所有LS项目的文档
                results = collection.query.fetch_objects(
                    filters=wvc.query.Filter.by_property("project").equal("LS"),
                    limit=10000
                )
                
                # 删除这些文档
                deleted_count = 0
                for obj in results.objects:
                    try:
                        collection.data.delete_by_id(obj.uuid)
                        deleted_count += 1
                    except Exception as e:
                        logger.warning(f"⚠️ 删除文档失败 {obj.uuid}: {e}")
                
                logger.info(f"✅ 删除了 {deleted_count} 个需求文档相关记录")
            
        except Exception as e:
            logger.error(f"❌ 清空需求文档数据失败: {e}")
            # 不抛出异常，继续执行
    
    def _docx_to_markdown(self, doc, file_path: Path) -> str:
        """将DOCX转换为Markdown格式"""
        markdown_content = []
        
        try:
            # 处理段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检测标题
                if self._is_heading(paragraph, text):
                    # 根据标题级别添加markdown标记
                    level = self._get_heading_level(paragraph, text)
                    markdown_content.append(f"{'#' * level} {text}")
                else:
                    # 普通段落
                    markdown_content.append(text)
            
            return "\n\n".join(markdown_content)
            
        except Exception as e:
            logger.error(f"DOCX转Markdown失败 {file_path}: {e}")
            return ""
    
    def _is_heading(self, paragraph, text: str) -> bool:
        """判断是否为标题"""
        # 检查段落样式
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # 检查文本格式（数字开头的标题）
        if re.match(r'^\d+\.?\d*\s+', text):
            return True
        
        # 检查是否为短文本且可能是标题
        if len(text) < 100 and any(keyword in text for keyword in [
            '系统', '模块', '功能', '接口', '设计', '架构', '需求', '背景', '目标', '范围',
            '流程', '方案', '实现', '测试', '部署', '运维', '总结', '附录'
        ]):
            return True
        
        return False
    
    def _get_heading_level(self, paragraph, text: str) -> int:
        """获取标题级别"""
        if paragraph.style.name.startswith('Heading'):
            try:
                return int(paragraph.style.name.replace('Heading ', ''))
            except:
                return 1
        
        # 根据数字格式判断级别
        match = re.match(r'^(\d+)\.?(\d*)\s+', text)
        if match:
            if match.group(2):  # 有二级数字，如 3.1
                return 2
            else:  # 只有一级数字，如 3
                return 1
        
        return 1
    
    def _split_markdown_by_sections(self, markdown_content: str, file_path: Path) -> List[Dict[str, Any]]:
        """按标题分割Markdown内容"""
        sections = []
        lines = markdown_content.split('\n')
        
        current_section = file_path.stem  # 使用文件名作为默认标题
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测标题行
            if line.startswith('#'):
                # 保存之前的段落
                if current_content:
                    sections.append({
                        'section': current_section,
                        'content': '\n'.join(current_content)
                    })
                
                # 开始新段落
                current_section = line.lstrip('#').strip()
                current_content = []
            else:
                # 普通内容
                current_content.append(line)
        
        # 处理最后一个段落
        if current_content:
            sections.append({
                'section': current_section,
                'content': '\n'.join(current_content)
            })
        
        # 如果没有找到任何段落，创建一个默认段落
        if not sections:
            sections.append({
                'section': file_path.stem,
                'content': markdown_content
            })
        
        return sections
    
    def _process_docx_file(self, file_path: Path) -> List[Document]:
        """处理 DOCX 文件 - 返回LangChain Document对象"""
        documents = []
        
        try:
            logger.info(f"📄 处理文档: {file_path.name}")
            
            # 打开DOCX文件
            try:
                doc = docx.Document(file_path)
            except Exception as docx_error:
                logger.error(f"❌ 无法打开DOCX文件 {file_path}: {docx_error}")
                return documents
            
            # 将DOCX转换为Markdown
            markdown_content = self._docx_to_markdown(doc, file_path)
            
            if not markdown_content:
                logger.warning(f"⚠️ 文档内容为空: {file_path.name}")
                return documents
            
            # 按标题分割内容
            sections = self._split_markdown_by_sections(markdown_content, file_path)
            
            for i, section in enumerate(sections):
                if len(section['content'].strip()) < 20:  # 过滤太短的内容
                    continue
                    
                # 创建LangChain Document对象
                document = Document(
                    page_content=section['content'],
                    metadata={
                        # 移除 file_path 元数据，只保留 file_name
                        'file_name': file_path.name,
                        'project': 'LS',  # 固定为LS项目
                        'file_type': 'docx',
                        'source_type': 'requirement_doc',
                        'title': section['section'],  # 使用title字段匹配项目期望
                        'chunk_index': i,
                        'total_chunks': len(sections),
                        'created_at': datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')  # RFC3339格式
                    }
                )
                documents.append(document)
            
            logger.info(f"✅ 文档处理完成: {file_path.name} ({len(documents)} 个段落)")
            
        except Exception as e:
            logger.error(f"❌ DOCX 文件处理失败 {file_path}: {e}")
        
        return documents
    
    def _create_weaviate_schema(self):
        """创建 Weaviate 模式"""
        try:
            collection_name = "Document"
            
            # 检查集合是否已存在
            if self.weaviate_client.collections.exists(collection_name):
                logger.info(f"📋 集合 {collection_name} 已存在，跳过创建")
                return collection_name
            
            # 创建集合
            self.weaviate_client.collections.create(
                name=collection_name,
                vectorizer_config=wvc.config.Configure.Vectorizer.none(),
                properties=[
                    wvc.config.Property(name="content", data_type=wvc.config.DataType.TEXT),
                    # 移除 file_path 属性，只保留 file_name
                    wvc.config.Property(name="file_name", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="project", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="file_type", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="source_type", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="title", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="chunk_index", data_type=wvc.config.DataType.INT),
                    wvc.config.Property(name="total_chunks", data_type=wvc.config.DataType.INT),
                    wvc.config.Property(name="created_at", data_type=wvc.config.DataType.DATE)
                ]
            )
            
            logger.info(f"✅ Weaviate 集合创建成功: {collection_name}")
            return collection_name
            
        except Exception as e:
            logger.error(f"❌ Weaviate 集合创建失败: {e}")
            raise
    
    def _insert_documents_batch(self, documents: List[Document], collection_name: str):
        """批量插入文档到Weaviate"""
        try:
            collection = self.weaviate_client.collections.get(collection_name)
            
            # 准备批量插入的数据
            with collection.batch.dynamic() as batch:
                for doc in documents:
                    # 生成向量
                    vector = self.embeddings_model.encode(doc.page_content).tolist()
                    
                    # 准备元数据
                    properties = {
                        "content": doc.page_content,
                        **doc.metadata
                    }
                    
                    # 添加到批次
                    batch.add_object(
                        properties=properties,
                        vector=vector
                    )
            
        except Exception as e:
            logger.error(f"❌ 批量插入文档失败: {e}")
            raise
    
    def scan_requirements_docs(self) -> List[Path]:
        """扫描需求文档目录，获取所有DOCX文件"""
        files = []
        
        try:
            if not self.requirements_docs_path.exists():
                logger.error(f"❌ 需求文档目录不存在: {self.requirements_docs_path}")
                return files
            
            for file_path in self.requirements_docs_path.glob("*.docx"):
                if file_path.is_file() and not file_path.name.startswith('~'):  # 排除临时文件
                    files.append(file_path)
            
            logger.info(f"📁 扫描完成，找到 {len(files)} 个DOCX文件")
            
            # 显示文件列表
            for file_path in files:
                logger.info(f"   📄 {file_path.name}")
            
        except Exception as e:
            logger.error(f"❌ 扫描需求文档目录失败: {e}")
        
        return files
    
    def process_all_files(self, files: List[Path]) -> List[Document]:
        """处理所有文件"""
        all_documents = []
        
        for i, file_path in enumerate(files, 1):
            logger.info(f"🔄 处理文件 {i}/{len(files)}: {file_path.name}")
            
            try:
                documents = self._process_docx_file(file_path)
                all_documents.extend(documents)
                
            except Exception as e:
                logger.error(f"❌ 处理文件失败 {file_path}: {e}")
        
        logger.info(f"📊 文件处理完成，总计生成 {len(all_documents)} 个文档段落")
        return all_documents
    
    def initialize_requirements_docs(self):
        """初始化需求文档向量库"""
        try:
            logger.info("🚀 开始初始化需求文档向量库...")
            
            # 1. 清空现有的需求文档数据
            self._clear_weaviate_database()
            
            # 2. 扫描需求文档文件
            files = self.scan_requirements_docs()
            if not files:
                logger.warning("⚠️ 没有找到DOCX文件")
                return
            
            # 3. 创建 Weaviate 模式
            collection_name = self._create_weaviate_schema()
            
            # 4. 处理文件
            documents = self.process_all_files(files)
            if not documents:
                logger.warning("⚠️ 没有生成任何文档")
                return
            
            # 5. 批量插入到Weaviate
            logger.info("📤 开始批量插入文档到Weaviate...")
            batch_size = 50  # 较小的批次大小，确保稳定性
            
            for i in range(0, len(documents), batch_size):
                batch = documents[i:i + batch_size]
                try:
                    self._insert_documents_batch(batch, collection_name)
                    logger.info(f"✅ 批次 {i//batch_size + 1}/{(len(documents) + batch_size - 1)//batch_size} 插入成功 ({len(batch)} 个文档)")
                except Exception as e:
                    logger.error(f"❌ 批次 {i//batch_size + 1} 插入失败: {e}")
            
            logger.info("🎉 需求文档向量库初始化完成！")
            
            # 验证结果
            collection = self.weaviate_client.collections.get(collection_name)
            total_docs = collection.query.fetch_objects(limit=1).objects
            if total_docs:
                logger.info(f"📋 验证：数据库中现有文档数量 > 0")
            
        except Exception as e:
            logger.error(f"❌ 需求文档向量库初始化失败: {e}")
            raise
    
    def query_test(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """测试查询功能"""
        try:
            # 生成查询向量
            query_vector = self.embeddings_model.encode(query).tolist()
            
            # 使用Weaviate进行向量搜索
            collection = self.weaviate_client.collections.get("Document")
            
            results = collection.query.near_vector(
                near_vector=query_vector,
                limit=k,
                return_metadata=["distance"]
            )
            
            # 格式化结果
            formatted_results = []
            for obj in results.objects:
                formatted_results.append({
                    'content': obj.properties.get('content', ''),
                    # 移除 file_path，使用 file_name
                    'file_name': obj.properties.get('file_name', ''),
                    'title': obj.properties.get('title', ''),
                    'similarity_score': 1 - obj.metadata.distance,
                    'distance': obj.metadata.distance
                })
            
            return formatted_results
            
        except Exception as e:
            logger.error(f"❌ 查询测试失败: {e}")
            return []
    
    def close(self):
        """关闭连接"""
        if self.weaviate_client:
            self.weaviate_client.close()

def main():
    """主函数"""
    initializer = None
    
    try:
        # 创建初始化器
        initializer = RequirementsDocsInitializer()
        
        # 初始化需求文档向量库
        initializer.initialize_requirements_docs()
        
        # 测试查询
        print("\n🔍 测试查询...")
        test_queries = [
            "需求背景",
            "系统架构",
            "接口设计",
            "用户权限"
        ]
        
        for query in test_queries:
            print(f"\n查询: {query}")
            results = initializer.query_test(query, k=3)
            for i, result in enumerate(results, 1):
                # 更新为使用 file_name 而不是 file_path
                print(f"{i}. 文件: {result['file_name']}")
                print(f"   标题: {result['title']}")
                print(f"   相似度: {result['similarity_score']:.3f}")
                print(f"   内容预览: {result['content'][:100]}...")
        
    except Exception as e:
        print(f"❌ 执行失败: {e}")
        import traceback
        traceback.print_exc()
        
    finally:
        if initializer:
            initializer.close()

if __name__ == "__main__":
    main()