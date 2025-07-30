#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
优化版需求文档向量化初始化脚本
解决分段策略问题，采用混合分段策略：
1. 智能语义分段 + 长度控制
2. 内容质量过滤 
3. 适当重叠处理
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

# 文件处理相关
import docx
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Weaviate 和 LangChain
import weaviate
import weaviate.classes as wvc
from langchain.schema import Document

# 项目配置
try:
    from src.utils.weaviate_helper import get_weaviate_client
except ImportError:
    import sys
    project_root = os.path.dirname(os.path.abspath(__file__))
    sys.path.insert(0, project_root)
    from src.utils.weaviate_helper import get_weaviate_client

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class OptimizedRequirementsDocsInitializer:
    """优化版需求文档初始化器 - 混合分段策略"""
    
    def __init__(self, requirements_docs_path: str = "/Users/renyu/Documents/knowledge_base/链数_LS/需求文档"):
        """
        初始化需求文档处理器
        
        Args:
            requirements_docs_path: 需求文档目录路径
        """
        self.requirements_docs_path = Path(requirements_docs_path)
        self.weaviate_client = None
        self.embeddings_model = None
        self.text_splitter = None
        
        # 只处理DOCX文件
        self.supported_extensions = {'.docx'}
        
        # 分段策略配置
        self.config = {
            'min_chunk_length': 50,     # 最小段落长度
            'max_chunk_length': 800,    # 最大段落长度，超过则二次分割
            'overlap_length': 100,      # 重叠长度
            'merge_threshold': 150,     # 短段落合并阈值
        }
        
        # 初始化组件
        self._initialize_components()
    
    def _initialize_components(self):
        """初始化各种组件"""
        try:
            # 初始化 Weaviate 客户端
            self.weaviate_client = get_weaviate_client()
            logger.info("✅ Weaviate 客户端初始化成功")
            
            # 初始化嵌入模型 - 使用bge-large-zh（1024维，中文优化）
            self.embeddings_model = SentenceTransformer('BAAI/bge-large-zh')
            logger.info("✅ 嵌入模型加载成功 (bge-large-zh, 1024维，中文优化)")
            
            # 初始化文本分割器（用于长段落的二次分割）
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.config['max_chunk_length'],
                chunk_overlap=self.config['overlap_length'],
                length_function=len,
                separators=["\n\n", "\n", "。", "；", "，", " ", ""]
            )
            logger.info("✅ 文本分割器初始化成功")
            
            logger.info("✅ 组件初始化完成")
            
        except Exception as e:
            logger.error(f"❌ 组件初始化失败: {e}")
            raise
    
    def _clear_weaviate_database(self):
        """清空Weaviate数据库中的需求文档数据"""
        try:
            logger.info("🗑️ 开始清空需求文档相关数据...")
            
            collection_name = "Document"
            if self.weaviate_client.collections.exists(collection_name):
                collection = self.weaviate_client.collections.get(collection_name)
                
                # 查询所有LS项目的文档
                results = collection.query.fetch_objects(
                    filters=wvc.query.Filter.by_property("project").equal("LS"),
                    limit=10000
                )
                
                # 删除这些文档
                deleted_count = 0
                for obj in results.objects:
                    try:
                        collection.data.delete_by_id(obj.uuid)
                        deleted_count += 1
                    except Exception as e:
                        logger.warning(f"⚠️ 删除文档失败 {obj.uuid}: {e}")
                
                logger.info(f"✅ 删除了 {deleted_count} 个需求文档相关记录")
            
        except Exception as e:
            logger.error(f"❌ 清空需求文档数据失败: {e}")
    
    def _is_meaningful_content(self, text: str) -> bool:
        """判断内容是否有意义（过滤格式化内容）"""
        text = text.strip()
        
        # 过滤条件
        meaningless_patterns = [
            r'^文档控制$',
            r'^中企云链（北京）金融信息服务有限公司$',
            r'^产品需求说明书$',
            r'^需求文档$',
            r'^\d+$',  # 纯数字
            r'^[一二三四五六七八九十]+$',  # 纯中文数字
            r'^第[一二三四五六七八九十\d]+章$',  # 章节标记
            r'^附录\s*[A-Z]?$',
            r'^目录$',
            r'^索引$',
        ]
        
        for pattern in meaningless_patterns:
            if re.match(pattern, text):
                return False
        
        # 内容太短且没有实质内容
        if len(text) < 20 and not any(keyword in text for keyword in [
            '功能', '接口', '需求', '设计', '实现', '调整', '新增', '修改', '删除'
        ]):
            return False
        
        return True
    
    def _is_heading(self, paragraph, text: str) -> bool:
        """改进的标题识别"""
        # 检查段落样式
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # 检查数字格式标题（更严格）
        if re.match(r'^\d+\.?\d*\s+.{2,}', text):  # 至少2个字符的标题内容
            return True
        
        # 检查是否为有意义的标题
        if len(text) < 100 and self._is_meaningful_content(text) and any(keyword in text for keyword in [
            '系统', '模块', '功能', '接口', '设计', '架构', '需求', '背景', '目标', '范围',
            '流程', '方案', '实现', '测试', '部署', '运维', '总结', '附录', '调整', '说明'
        ]):
            return True
        
        return False
    
    def _smart_segment_content(self, markdown_content: str, file_path: Path) -> List[Dict[str, Any]]:
        """智能分段策略"""
        sections = []
        lines = markdown_content.split('\n')
        
        current_section = file_path.stem
        current_content = []
        pending_short_sections = []  # 存储待合并的短段落
        
        def process_section(section_title: str, section_content: List[str]) -> List[Dict[str, Any]]:
            """处理单个段落"""
            content_text = '\n'.join(section_content).strip()
            
            # 过滤无意义内容
            if not self._is_meaningful_content(content_text):
                return []
            
            results = []
            content_length = len(content_text)
            
            if content_length < self.config['min_chunk_length']:
                # 太短的段落暂存，等待合并
                return [{'section': section_title, 'content': content_text, 'need_merge': True}]
            
            elif content_length > self.config['max_chunk_length']:
                # 太长的段落进行二次分割
                logger.info(f"📏 段落过长({content_length}字符)，进行二次分割: {section_title}")
                
                # 使用RecursiveCharacterTextSplitter分割
                text_chunks = self.text_splitter.split_text(content_text)
                
                for i, chunk in enumerate(text_chunks):
                    if len(chunk.strip()) >= self.config['min_chunk_length']:
                        chunk_title = f"{section_title}" if len(text_chunks) == 1 else f"{section_title}(第{i+1}部分)"
                        results.append({
                            'section': chunk_title,
                            'content': chunk.strip(),
                            'need_merge': False
                        })
                
                return results
            
            else:
                # 长度适中的段落直接使用
                return [{'section': section_title, 'content': content_text, 'need_merge': False}]
        
        # 第一遍：按标题分割
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测标题行
            if line.startswith('#'):
                # 处理之前的段落
                if current_content:
                    section_results = process_section(current_section, current_content)
                    sections.extend(section_results)
                
                # 开始新段落
                current_section = line.lstrip('#').strip()
                current_content = []
            else:
                current_content.append(line)
        
        # 处理最后一个段落
        if current_content:
            section_results = process_section(current_section, current_content)
            sections.extend(section_results)
        
        # 第二遍：处理需要合并的短段落
        final_sections = []
        accumulated_content = []
        accumulated_titles = []
        
        for section in sections:
            if section.get('need_merge', False):
                # 累积短段落
                accumulated_content.append(section['content'])
                accumulated_titles.append(section['section'])
                
                # 检查是否达到合并阈值
                total_length = sum(len(content) for content in accumulated_content)
                if total_length >= self.config['merge_threshold']:
                    # 合并并输出
                    merged_title = f"{accumulated_titles[0]}等{len(accumulated_titles)}项" if len(accumulated_titles) > 1 else accumulated_titles[0]
                    merged_content = '\n\n'.join(accumulated_content)
                    
                    final_sections.append({
                        'section': merged_title,
                        'content': merged_content
                    })
                    
                    # 清空累积
                    accumulated_content = []
                    accumulated_titles = []
            else:
                # 先输出累积的短段落（如果有）
                if accumulated_content:
                    merged_title = f"{accumulated_titles[0]}等{len(accumulated_titles)}项" if len(accumulated_titles) > 1 else accumulated_titles[0]
                    merged_content = '\n\n'.join(accumulated_content)
                    
                    final_sections.append({
                        'section': merged_title,
                        'content': merged_content
                    })
                    
                    accumulated_content = []
                    accumulated_titles = []
                
                # 输出当前段落
                final_sections.append({
                    'section': section['section'],
                    'content': section['content']
                })
        
        # 处理最后剩余的短段落
        if accumulated_content:
            merged_title = f"{accumulated_titles[0]}等{len(accumulated_titles)}项" if len(accumulated_titles) > 1 else accumulated_titles[0]
            merged_content = '\n\n'.join(accumulated_content)
            
            final_sections.append({
                'section': merged_title,
                'content': merged_content
            })
        
        # 如果没有任何有效段落，创建一个默认段落
        if not final_sections:
            if self._is_meaningful_content(markdown_content):
                final_sections.append({
                    'section': file_path.stem,
                    'content': markdown_content
                })
        
        return final_sections
    
    def _docx_to_markdown(self, doc, file_path: Path) -> str:
        """将DOCX转换为Markdown格式"""
        markdown_content = []
        
        try:
            # 处理段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检测标题
                if self._is_heading(paragraph, text):
                    # 根据标题级别添加markdown标记
                    level = self._get_heading_level(paragraph, text)
                    markdown_content.append(f"{'#' * level} {text}")
                else:
                    # 普通段落
                    markdown_content.append(text)
            
            return "\n\n".join(markdown_content)
            
        except Exception as e:
            logger.error(f"DOCX转Markdown失败 {file_path}: {e}")
            return ""
    
    def _get_heading_level(self, paragraph, text: str) -> int:
        """获取标题级别"""
        if paragraph.style.name.startswith('Heading'):
            try:
                return int(paragraph.style.name.replace('Heading ', ''))
            except:
                return 1
        
        # 根据数字格式判断级别
        match = re.match(r'^(\d+)\.?(\d*)\s+', text)
        if match:
            if match.group(2):  # 有二级数字，如 3.1
                return 2
            else:  # 只有一级数字，如 3
                return 1
        
        return 1
    
    def _extract_text_from_damaged_docx(self, file_path: Path) -> str:
        """从损坏的大DOCX文件中提取文本内容"""
        import zipfile
        import xml.etree.ElementTree as ET
        
        try:
            text_content = []
            
            logger.info(f"🔧 尝试使用zipfile方法处理损坏的大文档: {file_path.name}")
            
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # 列出所有文件
                file_list = docx_zip.namelist()
                
                # 尝试读取主文档内容
                main_doc_files = [
                    'word/document.xml',
                    'word/document2.xml', 
                    'document.xml'
                ]
                
                for doc_file in main_doc_files:
                    if doc_file in file_list:
                        try:
                            with docx_zip.open(doc_file) as xml_file:
                                xml_content = xml_file.read()
                                
                                # 解析XML并提取文本
                                root = ET.fromstring(xml_content)
                                
                                # 定义命名空间
                                namespaces = {
                                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                }
                                
                                # 提取所有文本节点
                                for text_elem in root.findall('.//w:t', namespaces):
                                    if text_elem.text:
                                        text_content.append(text_elem.text)
                                
                                logger.info(f"✅ 成功从 {doc_file} 提取文本")
                                break
                                
                        except Exception as e:
                            logger.warning(f"⚠️ 无法处理 {doc_file}: {e}")
                            continue
                
                # 如果主文档失败，尝试其他可能的文本文件
                if not text_content:
                    for file_name in file_list:
                        if file_name.endswith('.xml') and 'word' in file_name:
                            try:
                                with docx_zip.open(file_name) as xml_file:
                                    xml_content = xml_file.read().decode('utf-8', errors='ignore')
                                    # 简单的正则提取文本
                                    import re
                                    text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                                    if text_matches:
                                        text_content.extend(text_matches)
                                        logger.info(f"✅ 通过正则从 {file_name} 提取文本")
                                        break
                            except Exception as e:
                                continue
            
            if text_content:
                # 清理和组织文本
                cleaned_text = []
                current_line = ""
                
                for text in text_content:
                    text = text.strip()
                    if not text:
                        continue
                    
                    # 检查是否是新段落的开始
                    if any(char in text for char in ['.', '。', '!', '！', '?', '？']) and len(current_line) > 20:
                        current_line += text
                        cleaned_text.append(current_line)
                        current_line = ""
                    else:
                        current_line += text + " "
                
                # 添加最后一行
                if current_line.strip():
                    cleaned_text.append(current_line.strip())
                
                final_text = f"# {file_path.name}\n\n" + "\n\n".join(cleaned_text)
                
                logger.info(f"📝 成功恢复大文档文本内容，共 {len(cleaned_text)} 段落")
                return final_text
            
        except Exception as e:
            logger.error(f"❌ 大文档zipfile方法也失败: {e}")
        
        return ""
    
    def _process_docx_file(self, file_path: Path) -> List[Document]:
        """处理 DOCX 文件 - 使用优化的分段策略，支持大文档处理"""
        documents = []
        
        try:
            # 检查文件大小，大文档需要特殊处理
            file_size = file_path.stat().st_size
            size_mb = file_size / (1024 * 1024)
            logger.info(f"📄 处理文档: {file_path.name} ({size_mb:.1f}MB)")
            
            # 大文档预警和内存优化
            if size_mb > 10:
                logger.warning(f"⚠️ 检测到大文档 ({size_mb:.1f}MB)，启用特殊处理模式")
                # 对于超大文档，增加处理超时和内存优化
                import gc
                gc.collect()  # 强制垃圾回收
            
            # 打开DOCX文件
            try:
                # 对于大文档，使用更严格的错误处理
                if size_mb > 15:
                    logger.info(f"📋 超大文档处理模式: {file_path.name}")
                    # 设置更长的超时时间
                    import signal
                    def timeout_handler(signum, frame):
                        raise TimeoutError("文档处理超时")
                    
                    signal.signal(signal.SIGALRM, timeout_handler)
                    signal.alarm(300)  # 5分钟超时
                    
                    try:
                        doc = docx.Document(file_path)
                        signal.alarm(0)  # 取消超时
                    except TimeoutError:
                        logger.error(f"❌ 文档处理超时: {file_path.name}")
                        return documents
                    except Exception as e:
                        signal.alarm(0)
                        raise e
                else:
                    doc = docx.Document(file_path)
                    
            except Exception as docx_error:
                logger.error(f"❌ 无法打开DOCX文件 {file_path}: {docx_error}")
                # 对于损坏的大文档，尝试使用zipfile方法恢复
                if size_mb > 5:
                    logger.info(f"🔧 尝试使用zipfile方法恢复大文档...")
                    try:
                        recovered_content = self._extract_text_from_damaged_docx(file_path)
                        if recovered_content:
                            logger.info(f"✅ 大文档文本恢复成功: {file_path.name}")
                            # 创建恢复文档
                            doc_obj = Document(
                                page_content=recovered_content,
                                metadata={
                                    'file_name': file_path.name,
                                    'project': 'LS',
                                    'file_type': 'docx',
                                    'source_type': 'requirement_doc',
                                    'title': f'{file_path.stem}_恢复内容',
                                    'chunk_index': 0,
                                    'total_chunks': 1,
                                    'created_at': datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'),
                                    'processing_strategy': 'recovered_large_doc',
                                    'file_size_mb': f"{size_mb:.1f}MB"
                                }
                            )
                            return [doc_obj]
                    except Exception as recovery_error:
                        logger.error(f"❌ 大文档恢复也失败: {recovery_error}")
                
                return documents
            
            # 将DOCX转换为Markdown
            try:
                markdown_content = self._docx_to_markdown(doc, file_path)
                
                # 大文档内存优化：及时释放doc对象
                if size_mb > 10:
                    del doc
                    import gc
                    gc.collect()
                    
            except Exception as markdown_error:
                logger.error(f"❌ Markdown转换失败 {file_path}: {markdown_error}")
                if size_mb > 10:
                    logger.info(f"🔧 大文档转换失败，尝试简化处理...")
                    # 对于大文档转换失败，尝试直接读取段落文本
                    try:
                        simple_content = []
                        for paragraph in doc.paragraphs:
                            text = paragraph.text.strip()
                            if text and len(text) > 10:  # 过滤短内容
                                simple_content.append(text)
                        
                        if simple_content:
                            markdown_content = f"# {file_path.stem}\n\n" + "\n\n".join(simple_content)
                            logger.info(f"✅ 大文档简化处理成功，提取{len(simple_content)}个段落")
                        else:
                            return documents
                    except Exception as e:
                        logger.error(f"❌ 大文档简化处理也失败: {e}")
                        return documents
                else:
                    return documents
            
            if not markdown_content:
                logger.warning(f"⚠️ 文档内容为空: {file_path.name}")
                return documents
            
            # 对于超大文档，限制内容长度避免内存问题
            if size_mb > 15 and len(markdown_content) > 100000:
                logger.warning(f"⚠️ 超大文档内容截断: {file_path.name} (原长度: {len(markdown_content)})")
                markdown_content = markdown_content[:100000] + "\n\n[内容因文档过大被截断]"
            
            # 使用智能分段策略
            try:
                sections = self._smart_segment_content(markdown_content, file_path)
                
                # 大文档分段数量限制，避免产生过多向量
                if size_mb > 10 and len(sections) > 100:
                    logger.warning(f"⚠️ 大文档分段数量限制: {file_path.name} (原分段数: {len(sections)}, 限制为前100个)")
                    sections = sections[:100]
                    
            except Exception as segment_error:
                logger.error(f"❌ 智能分段失败 {file_path}: {segment_error}")
                # 降级为简单分段
                sections = [{'section': file_path.stem, 'content': markdown_content[:5000]}]
            
            for i, section in enumerate(sections):
                # 创建LangChain Document对象
                document = Document(
                    page_content=section['content'],
                    metadata={
                        'file_name': file_path.name,
                        'project': 'LS',
                        'file_type': 'docx',
                        'source_type': 'requirement_doc',
                        'title': section['section'],
                        'chunk_index': i,
                        'total_chunks': len(sections),
                        'created_at': datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'),
                        'processing_strategy': 'optimized_hybrid',
                        'file_size_mb': f"{size_mb:.1f}MB"
                    }
                )
                documents.append(document)
            
            logger.info(f"✅ 文档处理完成: {file_path.name} ({len(documents)} 个段落)")
            
            # 显示段落长度分布
            if documents:
                lengths = [len(doc.page_content) for doc in documents]
                logger.info(f"   📏 段落长度: 最短{min(lengths)}, 最长{max(lengths)}, 平均{sum(lengths)//len(lengths)}")
                
                # 大文档额外统计
                if size_mb > 5:
                    short_count = sum(1 for l in lengths if l < 100)
                    long_count = sum(1 for l in lengths if l > 1000)
                    logger.info(f"   📊 段落分布: 短段落(<100字符):{short_count}, 长段落(>1000字符):{long_count}")
            
        except Exception as e:
            logger.error(f"❌ DOCX 文件处理失败 {file_path}: {e}")
        
        return documents
    
    def _create_weaviate_schema(self):
        """创建 Weaviate 模式"""
        try:
            collection_name = "Document"
            
            # 检查集合是否已存在
            if self.weaviate_client.collections.exists(collection_name):
                logger.info(f"📋 集合 {collection_name} 已存在，跳过创建")
                return collection_name
            
            # 创建集合
            self.weaviate_client.collections.create(
                name=collection_name,
                vectorizer_config=wvc.config.Configure.Vectorizer.none(),
                properties=[
                    wvc.config.Property(name="content", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="file_name", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="project", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="file_type", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="source_type", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="title", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="chunk_index", data_type=wvc.config.DataType.INT),
                    wvc.config.Property(name="total_chunks", data_type=wvc.config.DataType.INT),
                    wvc.config.Property(name="created_at", data_type=wvc.config.DataType.DATE),
                    wvc.config.Property(name="processing_strategy", data_type=wvc.config.DataType.TEXT)
                ]
            )
            
            logger.info(f"✅ Weaviate 集合创建成功: {collection_name}")
            return collection_name
            
        except Exception as e:
            logger.error(f"❌ Weaviate 集合创建失败: {e}")
            raise
    
    def _insert_documents_batch(self, documents: List[Document], collection_name: str):
        """批量插入文档到Weaviate"""
        try:
            collection = self.weaviate_client.collections.get(collection_name)
            
            # 准备批量插入的数据
            with collection.batch.dynamic() as batch:
                for doc in documents:
                    # 生成向量
                    vector = self.embeddings_model.encode(doc.page_content).tolist()
                    
                    # 准备元数据
                    properties = {
                        "content": doc.page_content,
                        **doc.metadata
                    }
                    
                    # 添加到批次
                    batch.add_object(
                        properties=properties,
                        vector=vector
                    )
            
        except Exception as e:
            logger.error(f"❌ 批量插入文档失败: {e}")
            raise
    
    def scan_requirements_docs(self) -> List[Path]:
        """扫描需求文档目录，获取所有DOCX文件"""
        files = []
        
        try:
            if not self.requirements_docs_path.exists():
                logger.error(f"❌ 需求文档目录不存在: {self.requirements_docs_path}")
                return files
            
            for file_path in self.requirements_docs_path.glob("*.docx"):
                if file_path.is_file() and not file_path.name.startswith('~'):  # 排除临时文件
                    files.append(file_path)
            
            # 统计文档大小信息
            if files:
                total_size = sum(f.stat().st_size for f in files)
                large_files = [f for f in files if f.stat().st_size > 10*1024*1024]  # >10MB
                
                logger.info(f"📁 扫描完成，找到 {len(files)} 个DOCX文件")
                logger.info(f"📊 文档统计:")
                logger.info(f"   📑 文档总数: {len(files)} 个")
                logger.info(f"   📏 总大小: {total_size/(1024*1024):.1f}MB")
                logger.info(f"   ⚠️  大文档(>10MB): {len(large_files)} 个")
                
                if large_files:
                    logger.info(f"   📋 大文档列表:")
                    for lf in large_files:
                        size_mb = lf.stat().st_size / (1024*1024)
                        logger.info(f"      - {lf.name}: {size_mb:.1f}MB")
            else:
                logger.info(f"📁 扫描完成，未找到DOCX文件")
            
        except Exception as e:
            logger.error(f"❌ 扫描需求文档目录失败: {e}")
        
        return files
    
    def initialize_optimized_requirements_docs(self):
        """使用优化策略初始化需求文档向量库"""
        try:
            logger.info("🚀 开始使用优化策略初始化需求文档向量库...")
            logger.info(f"📋 分段策略配置: {self.config}")
            
            # 1. 清空现有的需求文档数据
            self._clear_weaviate_database()
            
            # 2. 扫描需求文档文件
            files = self.scan_requirements_docs()
            if not files:
                logger.warning("⚠️ 没有找到DOCX文件")
                return
            
            # 3. 创建 Weaviate 模式
            collection_name = self._create_weaviate_schema()
            
            # 4. 处理文件
            all_documents = []
            successful_files = 0
            failed_files = 0
            
            logger.info(f"🔄 开始处理 {len(files)} 个文档文件...")
            logger.info("=" * 60)
            
            for i, file_path in enumerate(files, 1):
                file_size_mb = file_path.stat().st_size / (1024*1024)
                logger.info(f"📄 [{i}/{len(files)}] 开始处理: {file_path.name} ({file_size_mb:.1f}MB)")
                
                try:
                    documents = self._process_docx_file(file_path)
                    if documents:
                        all_documents.extend(documents)
                        successful_files += 1
                        logger.info(f"✅ [{i}/{len(files)}] 处理成功: {file_path.name} -> {len(documents)} 个段落")
                    else:
                        failed_files += 1
                        logger.warning(f"⚠️ [{i}/{len(files)}] 处理结果为空: {file_path.name}")
                    
                except Exception as e:
                    failed_files += 1
                    logger.error(f"❌ [{i}/{len(files)}] 处理失败: {file_path.name} - {e}")
                
                # 显示进度统计
                logger.info(f"📊 当前进度: 成功{successful_files}个, 失败{failed_files}个, 累计段落{len(all_documents)}个")
                logger.info("-" * 60)
            
            # 处理完成统计
            logger.info("🎯 文档处理完成！")
            logger.info(f"📊 处理统计:")
            logger.info(f"   ✅ 成功处理: {successful_files} 个文件")
            logger.info(f"   ❌ 处理失败: {failed_files} 个文件")
            logger.info(f"   📑 生成段落: {len(all_documents)} 个")
            logger.info(f"   📈 成功率: {successful_files/(successful_files+failed_files)*100:.1f}%")
            logger.info("=" * 60)
            
            if not all_documents:
                logger.warning("⚠️ 没有生成任何文档")
                return
            
            # 5. 批量插入到Weaviate
            logger.info(f"📤 开始批量插入文档到Weaviate，总计 {len(all_documents)} 个段落...")
            batch_size = 50
            
            for i in range(0, len(all_documents), batch_size):
                batch = all_documents[i:i + batch_size]
                try:
                    self._insert_documents_batch(batch, collection_name)
                    logger.info(f"✅ 批次 {i//batch_size + 1}/{(len(all_documents) + batch_size - 1)//batch_size} 插入成功 ({len(batch)} 个文档)")
                except Exception as e:
                    logger.error(f"❌ 批次 {i//batch_size + 1} 插入失败: {e}")
            
            # 6. 统计分析
            lengths = [len(doc.page_content) for doc in all_documents]
            logger.info("📊 最终统计:")
            logger.info(f"   📑 总文档数: {len(all_documents)}")
            logger.info(f"   📏 平均长度: {sum(lengths)//len(lengths)} 字符")
            logger.info(f"   📏 长度范围: {min(lengths)} - {max(lengths)} 字符")
            
            # 长度分布
            short_count = sum(1 for l in lengths if l < 100)
            medium_count = sum(1 for l in lengths if 100 <= l < 400)
            long_count = sum(1 for l in lengths if l >= 400)
            
            logger.info(f"   📋 长度分布:")
            logger.info(f"      短段落(<100字符): {short_count} 个 ({short_count/len(lengths)*100:.1f}%)")
            logger.info(f"      中段落(100-399字符): {medium_count} 个 ({medium_count/len(lengths)*100:.1f}%)")  
            logger.info(f"      长段落(≥400字符): {long_count} 个 ({long_count/len(lengths)*100:.1f}%)")
            
            logger.info("🎉 优化版需求文档向量库初始化完成！")
            
        except Exception as e:
            logger.error(f"❌ 需求文档向量库初始化失败: {e}")
            raise
    
    def close(self):
        """关闭连接"""
        if self.weaviate_client:
            self.weaviate_client.close()

def main():
    """主函数"""
    initializer = None
    
    try:
        # 创建优化版初始化器
        initializer = OptimizedRequirementsDocsInitializer()
        
        # 使用优化策略初始化需求文档向量库
        initializer.initialize_optimized_requirements_docs()
        
    except Exception as e:
        print(f"❌ 执行失败: {e}")
        import traceback
        traceback.print_exc()
        
    finally:
        if initializer:
            initializer.close()

if __name__ == "__main__":
    main()