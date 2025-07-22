#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强分析器
提供高级文本分析、结构解析、内容理解等功能
"""

import os
import re
import json
import logging
import base64
from io import BytesIO
from typing import Dict, List, Any, Optional, Tuple
from collections import Counter
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

logger = logging.getLogger(__name__)

class EnhancedAnalyzer:
    """增强分析器 - 提供高级文本分析功能"""
    
    def __init__(self, base_url: str = "http://localhost:8082"):
        """初始化增强分析器"""
        self.text_patterns = self._init_text_patterns()
        self.structure_patterns = self._init_structure_patterns()
        self.base_url = base_url
        logger.info("增强分析器初始化完成")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx未安装，Word文档解析功能受限")
        if not PDF_AVAILABLE:
            logger.warning("pypdf未安装，PDF文档解析功能受限")
    
    def _init_text_patterns(self) -> Dict[str, re.Pattern]:
        """初始化文本模式"""
        return {
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'url': re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'),
            'phone': re.compile(r'(\d{3}[-.]?\d{3}[-.]?\d{4}|\d{3}[-.]?\d{4}[-.]?\d{4})'),
            'date': re.compile(r'\d{4}[-/]\d{1,2}[-/]\d{1,2}|\d{1,2}[-/]\d{1,2}[-/]\d{4}'),
            'time': re.compile(r'\d{1,2}:\d{2}(:\d{2})?(\s?[AaPp][Mm])?'),
            'code_block': re.compile(r'```[\s\S]*?```|`[^`]+`'),
            'heading': re.compile(r'^#{1,6}\s+.+$', re.MULTILINE),
            'list_item': re.compile(r'^\s*[-*+]\s+.+$|^\s*\d+\.\s+.+$', re.MULTILINE),
            'table_row': re.compile(r'\|.*\|'),
            'bold_text': re.compile(r'\*\*([^*]+)\*\*|__([^_]+)__'),
            'italic_text': re.compile(r'\*([^*]+)\*|_([^_]+)_'),
            'link': re.compile(r'\[([^\]]+)\]\(([^)]+)\)'),
        }
    
    def _init_structure_patterns(self) -> Dict[str, re.Pattern]:
        """初始化结构模式"""
        return {
            'section_header': re.compile(r'^#{1,6}\s+(.+)$', re.MULTILINE),
            'paragraph_break': re.compile(r'\n\s*\n'),
            'sentence_end': re.compile(r'[.!?]+\s+'),
            'word_boundary': re.compile(r'\b\w+\b'),
            'whitespace': re.compile(r'\s+'),
            'line_break': re.compile(r'\n'),
        }
    
    def analyze_text_comprehensive(self, text: str) -> Dict[str, Any]:
        """
        综合文本分析
        
        Args:
            text: 输入文本
            
        Returns:
            包含各种分析结果的字典
        """
        try:
            logger.info(f"开始综合文本分析，文本长度: {len(text)}")
            
            analysis_result = {
                'basic_stats': self._analyze_basic_stats(text),
                'structure_analysis': self._analyze_structure(text),
                'content_extraction': self._extract_content_elements(text),
                'linguistic_analysis': self._analyze_linguistics(text),
                'metadata': self._extract_metadata(text),
                'quality_metrics': self._analyze_quality(text),
                'summary': self._generate_summary(text)
            }
            
            logger.info("综合文本分析完成")
            return analysis_result
            
        except Exception as e:
            logger.error(f"综合文本分析失败: {e}")
            return self._get_fallback_analysis(text)
    
    def _analyze_basic_stats(self, text: str) -> Dict[str, Any]:
        """分析基础统计信息"""
        lines = text.split('\n')
        words = self.structure_patterns['word_boundary'].findall(text)
        sentences = self.structure_patterns['sentence_end'].split(text)
        paragraphs = self.structure_patterns['paragraph_break'].split(text.strip())
        
        return {
            'character_count': len(text),
            'character_count_no_spaces': len(text.replace(' ', '')),
            'word_count': len(words),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'paragraph_count': len([p for p in paragraphs if p.strip()]),
            'line_count': len(lines),
            'avg_words_per_sentence': len(words) / max(len(sentences), 1),
            'avg_sentences_per_paragraph': len(sentences) / max(len(paragraphs), 1),
            'reading_time_minutes': round(len(words) / 200, 1)  # 假设阅读速度200词/分钟
        }
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """分析文档结构"""
        headings = self.text_patterns['heading'].findall(text)
        lists = self.text_patterns['list_item'].findall(text)
        code_blocks = self.text_patterns['code_block'].findall(text)
        tables = self.text_patterns['table_row'].findall(text)
        
        # 分析标题层级
        heading_levels = []
        for heading in headings:
            level = len(heading) - len(heading.lstrip('#'))
            heading_levels.append(level)
        
        return {
            'has_headings': len(headings) > 0,
            'heading_count': len(headings),
            'heading_levels': list(set(heading_levels)) if heading_levels else [],
            'max_heading_depth': max(heading_levels) if heading_levels else 0,
            'has_lists': len(lists) > 0,
            'list_item_count': len(lists),
            'has_code': len(code_blocks) > 0,
            'code_block_count': len(code_blocks),
            'has_tables': len(tables) > 0,
            'table_row_count': len(tables),
            'structure_type': self._determine_structure_type(headings, lists, code_blocks, tables)
        }
    
    def _extract_content_elements(self, text: str) -> Dict[str, Any]:
        """提取内容元素"""
        emails = self.text_patterns['email'].findall(text)
        urls = self.text_patterns['url'].findall(text)
        phones = self.text_patterns['phone'].findall(text)
        dates = self.text_patterns['date'].findall(text)
        times = self.text_patterns['time'].findall(text)
        links = self.text_patterns['link'].findall(text)
        bold_text = self.text_patterns['bold_text'].findall(text)
        italic_text = self.text_patterns['italic_text'].findall(text)
        
        return {
            'contact_info': {
                'emails': list(set(emails)),
                'phones': list(set([phone[0] for phone in phones])),
                'has_contact_info': len(emails) > 0 or len(phones) > 0
            },
            'references': {
                'urls': list(set(urls)),
                'links': [{'text': link[0], 'url': link[1]} for link in links],
                'has_external_references': len(urls) > 0 or len(links) > 0
            },
            'temporal_info': {
                'dates': list(set(dates)),
                'times': list(set([time[0] for time in times])),
                'has_temporal_references': len(dates) > 0 or len(times) > 0
            },
            'formatting': {
                'bold_text': [text[0] or text[1] for text in bold_text],
                'italic_text': [text[0] or text[1] for text in italic_text],
                'has_formatting': len(bold_text) > 0 or len(italic_text) > 0
            }
        }
    
    def _analyze_linguistics(self, text: str) -> Dict[str, Any]:
        """语言学分析"""
        words = self.structure_patterns['word_boundary'].findall(text.lower())
        word_freq = Counter(words)
        
        # 简单的语言检测
        chinese_chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
        english_chars = sum(1 for char in text if char.isalpha() and ord(char) < 128)
        
        # 词汇多样性
        unique_words = len(set(words))
        total_words = len(words)
        lexical_diversity = unique_words / max(total_words, 1)
        
        return {
            'language': {
                'primary': "中文" if chinese_chars > english_chars else "英文",
                'chinese_char_count': chinese_chars,
                'english_char_count': english_chars,
                'is_multilingual': chinese_chars > 0 and english_chars > 0
            },
            'vocabulary': {
                'unique_words': unique_words,
                'total_words': total_words,
                'lexical_diversity': round(lexical_diversity, 3),
                'most_common_words': word_freq.most_common(10)
            },
            'complexity': {
                'avg_word_length': sum(len(word) for word in words) / max(len(words), 1),
                'complexity_score': self._calculate_complexity_score(text)
            }
        }
    
    def _extract_metadata(self, text: str) -> Dict[str, Any]:
        """提取元数据"""
        lines = text.split('\n')
        first_non_empty_line = next((line.strip() for line in lines if line.strip()), "")
        
        # 尝试识别标题
        title = ""
        if first_non_empty_line.startswith('#'):
            title = first_non_empty_line.lstrip('#').strip()
        elif first_non_empty_line:
            title = first_non_empty_line
        
        # 识别文档类型
        doc_type = self._identify_document_type(text)
        
        return {
            'title': title,
            'document_type': doc_type,
            'estimated_pages': max(1, len(text) // 2000),  # 估算页数
            'creation_date': datetime.now().isoformat(),
            'language_detected': "中文" if any('\u4e00' <= char <= '\u9fff' for char in text) else "英文",
            'content_summary': text[:200] + "..." if len(text) > 200 else text
        }
    
    def _analyze_quality(self, text: str) -> Dict[str, Any]:
        """分析文档质量"""
        words = self.structure_patterns['word_boundary'].findall(text)
        sentences = [s.strip() for s in self.structure_patterns['sentence_end'].split(text) if s.strip()]
        
        # 可读性评分
        avg_sentence_length = len(words) / max(len(sentences), 1)
        readability_score = max(1, min(10, 10 - (avg_sentence_length - 15) / 5))
        
        # 结构质量
        has_headings = bool(self.text_patterns['heading'].search(text))
        has_lists = bool(self.text_patterns['list_item'].search(text))
        structure_score = 5 + (2 if has_headings else 0) + (1 if has_lists else 0)
        
        # 内容完整性
        has_intro = len(text) > 200
        has_conclusion = "总结" in text or "结论" in text or "conclusion" in text.lower()
        completeness_score = 5 + (2 if has_intro else 0) + (1 if has_conclusion else 0)
        
        overall_score = (readability_score + structure_score + completeness_score) / 3
        
        return {
            'readability_score': round(readability_score, 1),
            'structure_score': round(structure_score, 1),
            'completeness_score': round(completeness_score, 1),
            'overall_quality_score': round(overall_score, 1),
            'quality_grade': self._get_quality_grade(overall_score),
            'recommendations': self._get_quality_recommendations(text)
        }
    
    def _generate_summary(self, text: str) -> Dict[str, Any]:
        """生成文档摘要"""
        sentences = [s.strip() for s in self.structure_patterns['sentence_end'].split(text) if s.strip()]
        
        # 简单的摘要生成（取前几句和关键句子）
        summary_sentences = sentences[:3] if len(sentences) >= 3 else sentences
        
        # 提取关键词
        words = self.structure_patterns['word_boundary'].findall(text.lower())
        word_freq = Counter(words)
        # 过滤常见停用词
        stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', 'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'a', 'an'}
        keywords = [word for word, freq in word_freq.most_common(20) if word not in stop_words and len(word) > 1]
        
        return {
            'brief_summary': ' '.join(summary_sentences),
            'key_points': sentences[:5] if len(sentences) >= 5 else sentences,
            'keywords': keywords[:10],
            'main_topics': self._extract_main_topics(text),
            'document_purpose': self._infer_document_purpose(text)
        }
    
    def _determine_structure_type(self, headings: List, lists: List, code_blocks: List, tables: List) -> str:
        """确定文档结构类型"""
        if headings and lists:
            return "hierarchical_structured"
        elif headings:
            return "hierarchical"
        elif lists:
            return "list_based"
        elif code_blocks:
            return "technical"
        elif tables:
            return "tabular"
        else:
            return "linear"
    
    def _calculate_complexity_score(self, text: str) -> float:
        """计算复杂度评分"""
        words = self.structure_patterns['word_boundary'].findall(text)
        sentences = self.structure_patterns['sentence_end'].split(text)
        
        avg_word_length = sum(len(word) for word in words) / max(len(words), 1)
        avg_sentence_length = len(words) / max(len(sentences), 1)
        
        # 简单的复杂度计算
        complexity = (avg_word_length * 0.3 + avg_sentence_length * 0.7) / 10
        return round(min(complexity, 1.0), 3)
    
    def _identify_document_type(self, text: str) -> str:
        """识别文档类型"""
        text_lower = text.lower()
        
        if "需求" in text or "requirement" in text_lower:
            return "requirements_document"
        elif "设计" in text or "design" in text_lower:
            return "design_document"
        elif "测试" in text or "test" in text_lower:
            return "test_document"
        elif "用户手册" in text or "manual" in text_lower:
            return "user_manual"
        elif "api" in text_lower or "接口" in text:
            return "api_documentation"
        elif "readme" in text_lower or "说明" in text:
            return "readme"
        else:
            return "general_document"
    
    def _get_quality_grade(self, score: float) -> str:
        """获取质量等级"""
        if score >= 8:
            return "Excellent"
        elif score >= 6:
            return "Good"
        elif score >= 4:
            return "Fair"
        else:
            return "Poor"
    
    def _get_quality_recommendations(self, text: str) -> List[str]:
        """获取质量改进建议"""
        recommendations = []
        
        if not self.text_patterns['heading'].search(text):
            recommendations.append("建议添加标题和章节结构")
        
        if not self.text_patterns['list_item'].search(text):
            recommendations.append("考虑使用列表来组织要点")
        
        words = len(self.structure_patterns['word_boundary'].findall(text))
        if words < 100:
            recommendations.append("内容可以更加详细和完整")
        
        sentences = self.structure_patterns['sentence_end'].split(text)
        avg_length = len(self.structure_patterns['word_boundary'].findall(text)) / max(len(sentences), 1)
        if avg_length > 25:
            recommendations.append("建议使用更短的句子提高可读性")
        
        return recommendations or ["文档质量良好"]
    
    def _extract_main_topics(self, text: str) -> List[str]:
        """提取主要主题"""
        headings = self.text_patterns['heading'].findall(text)
        
        if headings:
            # 从标题中提取主题
            topics = [heading.lstrip('#').strip() for heading in headings[:5]]
        else:
            # 从内容中提取关键短语
            words = self.structure_patterns['word_boundary'].findall(text.lower())
            word_freq = Counter(words)
            topics = [word for word, freq in word_freq.most_common(5) if len(word) > 2]
        
        return topics
    
    def _infer_document_purpose(self, text: str) -> str:
        """推断文档目的"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ["如何", "步骤", "教程", "指南", "how to", "tutorial", "guide"]):
            return "instructional"
        elif any(word in text_lower for word in ["分析", "研究", "报告", "analysis", "research", "report"]):
            return "analytical"
        elif any(word in text_lower for word in ["规范", "标准", "specification", "standard"]):
            return "specification"
        elif any(word in text_lower for word in ["介绍", "概述", "introduction", "overview"]):
            return "informational"
        else:
            return "general"
    
    def _get_fallback_analysis(self, text: str) -> Dict[str, Any]:
        """获取降级分析结果"""
        logger.warning("使用降级分析")
        
        return {
            'basic_stats': {
                'character_count': len(text),
                'word_count': len(text.split()),
                'line_count': len(text.split('\n'))
            },
            'structure_analysis': {
                'has_headings': '#' in text,
                'structure_type': 'unknown'
            },
            'content_extraction': {
                'contact_info': {'has_contact_info': False},
                'references': {'has_external_references': False},
                'temporal_info': {'has_temporal_references': False},
                'formatting': {'has_formatting': False}
            },
            'linguistic_analysis': {
                'language': {'primary': 'unknown'},
                'vocabulary': {'unique_words': 0, 'total_words': 0},
                'complexity': {'complexity_score': 0.5}
            },
            'metadata': {
                'title': 'unknown',
                'document_type': 'unknown',
                'language_detected': 'unknown'
            },
            'quality_metrics': {
                'overall_quality_score': 5.0,
                'quality_grade': 'Fair'
            },
            'summary': {
                'brief_summary': text[:200] + "..." if len(text) > 200 else text,
                'keywords': [],
                'main_topics': [],
                'document_purpose': 'unknown'
            }
        }

    def transform_file(self, file_content, file_name: str) -> Dict[str, Any]:
        """
        分析文件内容
        
        Args:
            file_content: 文件内容（可以是 bytes 或 str）
            file_name: 文件名
            
        Returns:
            分析结果字典
        """
        try:
            # 处理不同类型的输入
            logger.info(f"开始分析文件: {file_name}, 内容类型: {type(file_content)}, 大小: {len(file_content)}")
            
            # 统一转换为 bytes 类型以确保一致性
            if isinstance(file_content, str):
                logger.info("输入为字符串，转换为 bytes")
                file_content_bytes = file_content.encode('utf-8')
            elif isinstance(file_content, bytes):
                logger.info("输入为 bytes，直接使用")
                file_content_bytes = file_content
            else:
                raise TypeError(f"不支持的文件内容类型: {type(file_content)}")
            
            # 根据文件扩展名确定文件类型
            file_ext = file_name.split('.')[-1].lower() if '.' in file_name else ''
            
            if file_ext in ['doc', 'docx']:
                return self.parse_word_document(file_content_bytes, file_name)
            elif file_ext == 'pdf':
                return self.parse_pdf_document(file_content_bytes, file_name)
            elif file_ext in ['txt', 'md', 'markdown']:
                return self.parse_text_document(file_content_bytes, file_name)
            else:
                # 尝试作为文本处理
                return self.parse_text_document(file_content_bytes, file_name)
                
        except Exception as e:
            logger.error(f"文件分析失败: {e}")
            return self._get_error_result(file_name, str(e))
    
    def parse_word_document(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        解析Word文档
        
        Args:
            file_content: Word文档的二进制内容
            file_name: 文件名
            
        Returns:
            解析结果字典
        """
        if not DOCX_AVAILABLE:
            return self._get_fallback_result(file_name, "Word解析", "python-docx未安装")
        
        try:
            # 检查文件格式
            file_ext = file_name.split('.')[-1].lower() if '.' in file_name else ''
            
            # 检查是否为旧格式的.doc文件
            if file_ext == 'doc' or (file_content.startswith(b'\xd0\xcf\x11\xe0') or 
                                   file_content.startswith(b'\x50\x4b\x03\x04') == False):
                logger.warning(f"检测到旧格式.doc文件或非ZIP格式文件: {file_name}")
                return self._handle_legacy_doc_file(file_content, file_name)
            
            # 检查文件是否为有效的ZIP文件（.docx格式基于ZIP）
            try:
                import zipfile
                file_stream = BytesIO(file_content)
                with zipfile.ZipFile(file_stream, 'r') as zip_ref:
                    # 检查是否包含Word文档必需的文件
                    required_files = ['word/document.xml', '[Content_Types].xml']
                    zip_files = zip_ref.namelist()
                    if not any(req_file in zip_files for req_file in required_files):
                        logger.warning(f"文件结构不符合Word文档格式: {file_name}")
                        return self._handle_invalid_word_file(file_content, file_name)
            except zipfile.BadZipFile:
                logger.warning(f"文件不是有效的ZIP格式: {file_name}")
                return self._handle_non_zip_file(file_content, file_name)
            
            # 创建文件流
            file_stream = BytesIO(file_content)
            
            # 解析文档
            doc = Document(file_stream)
            
            # 转换为Markdown格式
            markdown_content = self._convert_word_to_markdown(doc, file_content, file_name)
    
            
            # 提取结构化信息用于统计
            paragraphs_info = []
            tables_info = []
            
            # 统计段落信息
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    paragraphs_info.append({
                        "index": i,
                        "text": paragraph.text,
                        "style": paragraph.style.name if paragraph.style else "Normal"
                    })
            
            # 统计表格信息
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    if any(row_data):
                        table_data.append(row_data)
                
                if table_data:
                    tables_info.append({
                        "index": table_idx,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0,
                        "data": table_data
                    })
            
            # 图片信息已在Markdown转换过程中处理
            has_images = "![" in markdown_content  # 简单检查是否有图片
            image_count = markdown_content.count("![")  # 统计图片数量
            
            
            # 构建结果
            result = {
                "text_content": markdown_content,  # 使用转换后的Markdown格式文本
                "file_type": "word",
                "file_name": file_name,
                "file_size": len(file_content),
                "analysis_method": "enhanced_word_to_markdown",
                "document_structure": {
                    "paragraphs": paragraphs_info,
                    "tables": tables_info,
                    "has_images": has_images,
                    "image_count": image_count,
                    "total_paragraphs": len(paragraphs_info),
                    "total_tables": len(tables_info),
                    "format": "markdown"
                },
                "extraction_summary": {
                    "paragraphs_extracted": len(paragraphs_info),
                    "tables_extracted": len(tables_info),
                    "images_detected": image_count,
                    "total_text_length": len(markdown_content),
                    "output_format": "markdown"
                }
            }
            
            logger.info(f"Word文档解析完成: {file_name}")
            return result
            
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            return self._get_error_result(file_name, f"Word解析失败: {str(e)}")
    
    def check_has_images(self, file_stream: BytesIO) -> Tuple[bool, int]:
        """
        检查Word文档是否包含图片
        
        Args:
            file_stream: 文件流
            
        Returns:
            (是否有图片, 图片数量)
        """
        try:
            # 重置流位置
            file_stream.seek(0)
            doc = Document(file_stream)
            
            image_count = 0
            
            # 检查文档中的所有关系（relationships）
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            # 另一种方法：检查内联形状
            try:
                from docx.oxml.ns import qn
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        for drawing in run._element.xpath('.//w:drawing'):
                            image_count += 1
            except:
                pass
            
            return image_count > 0, image_count
            
        except Exception as e:
            logger.warning(f"图片检测失败: {e}")
            return False, 0
    
    def parse_pdf_document(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        解析PDF文档
        
        Args:
            file_content: PDF文档的二进制内容
            file_name: 文件名
            
        Returns:
            解析结果字典
        """
        if not PDF_AVAILABLE:
            return self._get_fallback_result(file_name, "PDF解析", "pypdf未安装")
        
        try:
            file_stream = BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(file_stream)
            
            # 提取文本
            text_parts = []
            pages_info = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                        pages_info.append({
                            "page_number": page_num + 1,
                            "text_length": len(page_text),
                            "has_content": bool(page_text.strip())
                        })
                except Exception as e:
                    logger.warning(f"PDF页面 {page_num + 1} 解析失败: {e}")
                    pages_info.append({
                        "page_number": page_num + 1,
                        "text_length": 0,
                        "has_content": False,
                        "error": str(e)
                    })
            
            # 合并文本
            content = '\n'.join(text_parts)
            
            # 进行综合文本分析
            text_analysis = self.analyze_text_comprehensive(content)
            
            # 获取PDF元数据
            metadata = {}
            try:
                if pdf_reader.metadata:
                    metadata = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', ''),
                        "creation_date": str(pdf_reader.metadata.get('/CreationDate', '')),
                        "modification_date": str(pdf_reader.metadata.get('/ModDate', ''))
                    }
            except Exception as e:
                logger.warning(f"PDF元数据提取失败: {e}")
            
            result = {
                "text_content": content,
                "file_type": "pdf",
                "file_name": file_name,
                "file_size": len(file_content),
                "analysis_method": "enhanced_pdf_parsing",
                "document_structure": {
                    "total_pages": len(pdf_reader.pages),
                    "pages_with_content": len([p for p in pages_info if p["has_content"]]),
                    "pages_info": pages_info,
                    "metadata": metadata
                },
                "content_analysis": text_analysis,
                "extraction_summary": {
                    "pages_processed": len(pdf_reader.pages),
                    "pages_with_text": len([p for p in pages_info if p["has_content"]]),
                    "total_text_length": len(content)
                }
            }
            
            logger.info(f"PDF文档解析完成: {file_name}")
            return result
            
        except Exception as e:
            logger.error(f"PDF文档解析失败: {e}")
            return self._get_error_result(file_name, f"PDF解析失败: {str(e)}")
    
    def parse_text_document(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        解析文本文档
        
        Args:
            file_content: 文本文档的二进制内容
            file_name: 文件名
            
        Returns:
            解析结果字典
        """
        try:
            # 使用改进的编码检测
            content, used_encoding = self._detect_encoding_and_decode(file_content)
            
            # 进行综合文本分析
            text_analysis = self.analyze_text_comprehensive(content)
            
            # 分析文本格式
            file_ext = file_name.split('.')[-1].lower() if '.' in file_name else 'txt'
            format_analysis = self._analyze_text_format(content, file_ext)
            
            result = {
                "text_content": content,
                "file_type": "text",
                "file_name": file_name,
                "file_size": len(file_content),
                "analysis_method": "enhanced_text_parsing",
                "encoding_info": {
                    "detected_encoding": used_encoding
                },
                "format_analysis": format_analysis,
                "content_analysis": text_analysis,
                "extraction_summary": {
                    "total_text_length": len(content),
                    "encoding_used": used_encoding,
                    "format_detected": format_analysis.get("format_type", "plain_text")
                }
            }
            
            logger.info(f"文本文档解析完成: {file_name}")
            return result
            
        except Exception as e:
            logger.error(f"文本文档解析失败: {e}")
            return self._get_error_result(file_name, f"文本解析失败: {str(e)}")
    
    def _analyze_text_format(self, content: str, file_ext: str) -> Dict[str, Any]:
        """分析文本格式"""
        format_info = {
            "file_extension": file_ext,
            "format_type": "plain_text",
            "special_features": []
        }
        
        if file_ext in ['md', 'markdown']:
            format_info["format_type"] = "markdown"
            # 检查Markdown特征
            if re.search(r'^#{1,6}\s', content, re.MULTILINE):
                format_info["special_features"].append("headings")
            if re.search(r'^\s*[-*+]\s', content, re.MULTILINE):
                format_info["special_features"].append("bullet_lists")
            if re.search(r'^\s*\d+\.\s', content, re.MULTILINE):
                format_info["special_features"].append("numbered_lists")
            if re.search(r'```[\s\S]*?```', content):
                format_info["special_features"].append("code_blocks")
            if re.search(r'\[([^\]]+)\]\(([^)]+)\)', content):
                format_info["special_features"].append("links")
            if re.search(r'\|.*\|', content):
                format_info["special_features"].append("tables")
        
        elif file_ext in ['json']:
            format_info["format_type"] = "json"
            try:
                json.loads(content)
                format_info["special_features"].append("valid_json")
            except:
                format_info["special_features"].append("invalid_json")
        
        elif file_ext in ['xml', 'html', 'htm']:
            format_info["format_type"] = "markup"
            if re.search(r'<[^>]+>', content):
                format_info["special_features"].append("markup_tags")
        
        elif file_ext in ['csv']:
            format_info["format_type"] = "csv"
            lines = content.split('\n')
            if len(lines) > 1:
                # 检查是否有一致的分隔符
                separators = [',', ';', '\t', '|']
                for sep in separators:
                    if all(line.count(sep) == lines[0].count(sep) for line in lines[1:3] if line.strip()):
                        format_info["special_features"].append(f"separator_{sep}")
                        break
        
        return format_info

    def _get_error_result(self, file_name: str, error_message: str) -> Dict[str, Any]:
        """获取错误结果"""
        return {
            "text_content": f"文件解析失败: {error_message}",
            "file_type": "unknown",
            "file_name": file_name,
            "file_size": 0,
            "analysis_method": "error_fallback",
            "error": error_message,
            "content_analysis": self._get_fallback_analysis(""),
            "extraction_summary": {
                "success": False,
                "error": error_message
            }
        }
    
    def _get_fallback_result(self, file_name: str, analysis_type: str, reason: str) -> Dict[str, Any]:
        """获取降级结果"""
        return {
            "text_content": f"{analysis_type}功能不可用: {reason}",
            "file_type": "unsupported",
            "file_name": file_name,
            "file_size": 0,
            "analysis_method": "fallback",
            "reason": reason,
            "content_analysis": self._get_fallback_analysis(""),
            "extraction_summary": {
                "success": False,
                "fallback_reason": reason
            }
        }
    
    def _handle_legacy_doc_file(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """处理旧格式的.doc文件"""
        logger.info(f"尝试处理旧格式Word文档: {file_name}")
        
        try:
            # 尝试使用其他库处理.doc文件
            # 如果没有安装专门的.doc处理库，提供降级处理
            
            # 检查是否有python-docx2txt库
            try:
                import docx2txt
                # 虽然docx2txt主要用于.docx，但可以尝试
                text_content = docx2txt.process(BytesIO(file_content))
                if text_content.strip():
                    return self._create_legacy_doc_result(text_content, file_name, file_content)
            except ImportError:
                pass
            except Exception:
                pass
            
            # 降级：尝试提取部分文本内容
            # .doc文件是二进制格式，我们只能提取可能的文本片段
            text_fragments = self._extract_text_from_binary(file_content)
            
            return {
                "text_content": text_fragments,
                "file_type": "legacy_word",
                "file_name": file_name,
                "file_size": len(file_content),
                "analysis_method": "legacy_doc_fallback",
                "warning": "这是旧格式的.doc文件，建议转换为.docx格式以获得更好的解析效果",
                "content_analysis": self.analyze_text_comprehensive(text_fragments) if text_fragments.strip() else self._get_fallback_analysis(""),
                "extraction_summary": {
                    "success": True,
                    "method": "binary_text_extraction",
                    "note": "从二进制文件中提取的文本可能不完整"
                }
            }
            
        except Exception as e:
            logger.error(f"旧格式Word文档处理失败: {e}")
            return self._get_error_result(file_name, f"旧格式.doc文件处理失败: {str(e)}")
    
    def _handle_non_zip_file(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """处理非ZIP格式的文件"""
        logger.info(f"文件不是ZIP格式，尝试作为文本处理: {file_name}")
        
        try:
            # 尝试作为纯文本文件处理
            return self.parse_text_document(file_content, file_name)
        except Exception as e:
            logger.error(f"非ZIP文件处理失败: {e}")
            return self._get_error_result(file_name, f"文件格式不支持: {str(e)}")
    
    def _handle_invalid_word_file(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """处理无效的Word文件"""
        logger.info(f"文件结构不符合Word文档标准: {file_name}")
        
        try:
            # 尝试作为压缩文件解压并提取文本
            import zipfile
            file_stream = BytesIO(file_content)
            text_parts = []
            
            with zipfile.ZipFile(file_stream, 'r') as zip_ref:
                for file_info in zip_ref.filelist:
                    if file_info.filename.endswith('.xml') or file_info.filename.endswith('.txt'):
                        try:
                            with zip_ref.open(file_info.filename) as xml_file:
                                xml_bytes = xml_file.read()
                                # 使用智能编码检测
                                content, _ = self._detect_encoding_and_decode(xml_bytes)
                                # 简单清理XML标签
                                import re
                                clean_text = re.sub(r'<[^>]+>', ' ', content)
                                clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                                if clean_text:
                                    text_parts.append(clean_text)
                        except Exception:
                            continue
            
            extracted_text = '\n'.join(text_parts)
            
            return {
                "text_content": extracted_text,
                "file_type": "corrupted_word",
                "file_name": file_name,
                "file_size": len(file_content),
                "analysis_method": "xml_extraction_fallback",
                "warning": "文件可能已损坏或格式异常，提取的内容可能不完整",
                "content_analysis": self.analyze_text_comprehensive(extracted_text) if extracted_text.strip() else self._get_fallback_analysis(""),
                "extraction_summary": {
                    "success": True,
                    "method": "xml_text_extraction",
                    "files_processed": len(text_parts)
                }
            }
            
        except Exception as e:
            logger.error(f"损坏Word文件处理失败: {e}")
            return self._get_error_result(file_name, f"Word文件损坏或格式异常: {str(e)}")
    
    def _detect_encoding_and_decode(self, file_content: bytes) -> Tuple[str, str]:
        """
        智能检测文件编码并解码
        
        Args:
            file_content: 文件的二进制内容
            
        Returns:
            (解码后的文本内容, 使用的编码)
        """
        try:
            # 检查BOM（字节顺序标记）
            if file_content.startswith(b'\xef\xbb\xbf'):
                return file_content[3:].decode('utf-8'), 'utf-8-bom'
            elif file_content.startswith(b'\xff\xfe'):
                return file_content[2:].decode('utf-16le'), 'utf-16le'
            elif file_content.startswith(b'\xfe\xff'):
                return file_content[2:].decode('utf-16be'), 'utf-16be'
            elif file_content.startswith(b'\xff\xfe\x00\x00'):
                return file_content[4:].decode('utf-32le'), 'utf-32le'
            elif file_content.startswith(b'\x00\x00\xfe\xff'):
                return file_content[4:].decode('utf-32be'), 'utf-32be'
            
            # 尝试使用chardet进行自动检测（如果可用）
            try:
                import chardet
                detected = chardet.detect(file_content)
                if detected and detected['confidence'] > 0.7:
                    encoding = detected['encoding']
                    if encoding:
                        return file_content.decode(encoding), f"{encoding} (chardet detected)"
            except ImportError:
                pass
            except Exception as e:
                logger.warning(f"chardet检测失败: {e}")
            
            # 手动尝试常见的中文编码
            chinese_encodings = [
                'utf-8',
                'gbk',
                'gb2312',
                'gb18030',
                'big5',
                'cp936',  # Windows中文简体
                'cp950',  # Windows中文繁体
            ]
            
            # 优先测试中文编码
            for encoding in chinese_encodings:
                try:
                    decoded = file_content.decode(encoding)
                    # 检查是否包含中文字符且没有明显的乱码
                    chinese_chars = sum(1 for char in decoded if '\u4e00' <= char <= '\u9fff')
                    total_chars = len(decoded)
                    
                    # 如果包含中文字符或者解码成功且内容合理
                    if chinese_chars > 0 or (total_chars > 0 and self._is_valid_text(decoded)):
                        return decoded, encoding
                except UnicodeDecodeError:
                    continue
                except Exception:
                    continue
            
            # 尝试其他常见编码
            other_encodings = [
                'latin1',
                'iso-8859-1',
                'cp1252',  # Windows西欧
                'utf-16',
                'utf-16le',
                'utf-16be',
                'ascii',
            ]
            
            for encoding in other_encodings:
                try:
                    decoded = file_content.decode(encoding)
                    if self._is_valid_text(decoded):
                        return decoded, encoding
                except UnicodeDecodeError:
                    continue
                except Exception:
                    continue
            
            # 最后降级处理：使用UTF-8并忽略错误
            decoded = file_content.decode('utf-8', errors='ignore')
            return decoded, 'utf-8 (with errors ignored)'
            
        except Exception as e:
            logger.error(f"编码检测失败: {e}")
            # 最终降级
            try:
                return file_content.decode('utf-8', errors='replace'), 'utf-8 (with replacement)'
            except:
                return str(file_content), 'binary (converted to string)'
    
    def _is_valid_text(self, text: str) -> bool:
        """
        检查文本是否看起来像有效的文本内容
        
        Args:
            text: 要检查的文本
            
        Returns:
            是否为有效文本
        """
        if not text or len(text.strip()) == 0:
            return False
        
        # 检查控制字符比例
        control_chars = sum(1 for char in text if ord(char) < 32 and char not in '\n\r\t')
        if len(text) > 0 and control_chars / len(text) > 0.3:
            return False
        
        # 检查是否有可读的字符
        printable_chars = sum(1 for char in text if char.isprintable() or char in '\n\r\t')
        if len(text) > 0 and printable_chars / len(text) < 0.7:
            return False
        
        return True
    
    def _extract_text_from_binary(self, file_content: bytes) -> str:
        """从二进制文件中提取可能的文本内容"""
        try:
            # 使用改进的编码检测
            content, encoding = self._detect_encoding_and_decode(file_content)
            
            # 如果检测到的内容看起来不像文本，尝试提取文本片段
            if not self._is_valid_text(content) or 'with errors' in encoding or 'binary' in encoding:
                import re
                text_fragments = []
                
                # 尝试不同的编码提取文本片段
                encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'latin1', 'utf-16', 'utf-16le', 'utf-16be']
                
                for encoding in encodings:
                    try:
                        decoded = file_content.decode(encoding, errors='ignore')
                        # 查找连续的可打印字符串（长度大于3）
                        # 改进正则表达式以更好地识别中文文本
                        printable_strings = re.findall(r'[\u4e00-\u9fff\w\s.,!?;:，。！？；：]{4,}', decoded)
                        text_fragments.extend(printable_strings)
                    except:
                        continue
                
                # 去重并清理
                unique_fragments = list(set(text_fragments))
                # 过滤并排序片段
                clean_fragments = []
                for frag in unique_fragments:
                    cleaned = frag.strip()
                    if len(cleaned) > 3:
                        # 优先保留包含中文的片段
                        chinese_chars = sum(1 for char in cleaned if '\u4e00' <= char <= '\u9fff')
                        if chinese_chars > 0 or self._is_valid_text(cleaned):
                            clean_fragments.append(cleaned)
                
                # 按长度排序，保留较长的片段
                clean_fragments.sort(key=len, reverse=True)
                return '\n'.join(clean_fragments[:30])  # 限制数量避免过多内容
            
            return content
            
        except Exception as e:
            logger.warning(f"二进制文本提取失败: {e}")
            return "无法从文件中提取文本内容"
    
    def _convert_word_to_markdown(self, doc, file_content: bytes, file_name: str) -> str:
        """
        将Word文档转换为Markdown格式
        
        Args:
            doc: python-docx Document对象
            file_content: 原始文件内容
            file_name: 文件名
            
        Returns:
            Markdown格式的文本内容
        """
        markdown_parts = []
        
        # 设置图片保存目录
        try:
            temp_dir = os.path.join('uploads', 'temp')
        except:
            temp_dir = os.path.join('analydesign', 'uploads', 'temp')
        
        # 确保目录存在
        os.makedirs(temp_dir, exist_ok=True)
        
        # 提取并保存图片，建立关系映射
        image_relationships = self._extract_and_save_images(doc, temp_dir, file_name)
        
        # 处理文档内容
        table_index = 0
        
        for element in doc.element.body:
            element_tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            
            if element_tag == 'p':  # 段落
                paragraph = self._find_paragraph_by_element(doc, element)
                if paragraph:
                    markdown_paragraph = self._convert_paragraph_to_markdown(paragraph, image_relationships)
                    if markdown_paragraph.strip():
                        markdown_parts.append(markdown_paragraph)
                        
            elif element_tag == 'tbl':  # 表格
                if table_index < len(doc.tables):
                    table = doc.tables[table_index]
                    markdown_table = self._convert_table_to_markdown(table)
                    if markdown_table.strip():
                        markdown_parts.append(markdown_table)
                    table_index += 1
        
        # 如果上面的方法没有提取到内容，使用简化方法
        if not markdown_parts:
            logger.warning("使用简化的Markdown转换方法")
            markdown_parts = self._simple_word_to_markdown(doc, image_relationships)
        
        return '\n\n'.join(markdown_parts)
    
    def _extract_and_save_images(self, doc, temp_dir: str, file_name: str) -> Dict[str, str]:
        """
        提取并保存Word文档中的图片
        
        Returns:
            图片关系映射 {relationship_id: markdown_image_syntax}
        """
        image_relationships = {}
        image_count = 0
        
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        # 确定图片扩展名
                        if rel.target_ref.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                            ext = rel.target_ref.split('.')[-1]
                        else:
                            # 根据图片数据判断格式
                            if image_data.startswith(b'\x89PNG'):
                                ext = 'png'
                            elif image_data.startswith(b'\xff\xd8'):
                                ext = 'jpg'
                            elif image_data.startswith(b'GIF'):
                                ext = 'gif'
                            else:
                                ext = 'png'  # 默认
                        
                        # 生成唯一文件名
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        base_name = os.path.splitext(file_name)[0]

                        # 截取文件名中_前面的部分
                        base_name = base_name.split('_')[0]

                        
                        image_filename = f"{base_name}_img_{timestamp}_{image_count}.{ext}"
                        image_path = os.path.join(temp_dir, image_filename)
                        
                        # 保存图片
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)
                        
                        # 创建Markdown图片语法 - 使用完整URL
                        full_url = f"{self.base_url}/uploads/temp/{image_filename}"
                        markdown_image = f"![图片{image_count + 1}]({full_url})"
                        
                        # 建立关系映射
                        image_relationships[rel.rId] = markdown_image
                        
                        image_count += 1
                        logger.info(f"图片已保存: {full_url}")
                        
                    except Exception as e:
                        logger.warning(f"保存图片失败: {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"提取图片失败: {e}")
        
        return image_relationships
    
    def _find_paragraph_by_element(self, doc, element):
        """根据element查找对应的paragraph对象"""
        for paragraph in doc.paragraphs:
            if paragraph._element == element:
                return paragraph
        return None
    
    def _convert_paragraph_to_markdown(self, paragraph, image_relationships: Dict[str, str]) -> str:
        """将段落转换为Markdown格式"""
        text = paragraph.text.strip()
        
        if not text:
            # 检查是否包含图片
            for run in paragraph.runs:
                for drawing in run._element.xpath('.//w:drawing'):
                    # 查找图片引用
                    for blip in drawing.xpath('.//a:blip'):
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id and embed_id in image_relationships:
                            return image_relationships[embed_id]
            return ""
        
        # 检测标题级别
        style_name = paragraph.style.name if paragraph.style else "Normal"
        
        if "Heading" in style_name or "标题" in style_name:
            # 提取标题级别
            if "1" in style_name:
                return f"# {text}"
            elif "2" in style_name:
                return f"## {text}"
            elif "3" in style_name:
                return f"### {text}"
            elif "4" in style_name:
                return f"#### {text}"
            elif "5" in style_name:
                return f"##### {text}"
            else:
                return f"## {text}"  # 默认二级标题
        
        # 检查是否为列表项
        if paragraph._element.xpath('.//w:numPr') or paragraph._element.xpath('.//w:pPr/w:numPr'):
            return f"- {text}"
        
        # 处理格式化文本
        formatted_text = self._process_text_formatting(paragraph)
        
        # 处理段落中的图片
        for run in paragraph.runs:
            for drawing in run._element.xpath('.//w:drawing'):
                for blip in drawing.xpath('.//a:blip'):
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id and embed_id in image_relationships:
                        formatted_text += f"\n\n{image_relationships[embed_id]}"
        
        return formatted_text
    
    def _process_text_formatting(self, paragraph) -> str:
        """处理文本格式化（粗体、斜体等）"""
        result = ""
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
                
            # 应用格式
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            
            result += text
        
        return result
    
    def _convert_table_to_markdown(self, table) -> str:
        """将表格转换为Markdown格式"""
        if not table.rows:
            return ""
        
        markdown_rows = []
        
        # 处理表头（第一行）
        header_row = table.rows[0]
        header_cells = [cell.text.strip() for cell in header_row.cells]
        if any(header_cells):  # 确保表头不为空
            markdown_rows.append("| " + " | ".join(header_cells) + " |")
            markdown_rows.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        
        # 处理数据行
        for row in table.rows[1:] if len(table.rows) > 1 else table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # 只添加非空行
                markdown_rows.append("| " + " | ".join(cells) + " |")
        
        return "\n".join(markdown_rows)
    
    def _simple_word_to_markdown(self, doc, image_relationships: Dict[str, str]) -> List[str]:
        """简化的Word到Markdown转换方法"""
        markdown_parts = []
        
        # 处理段落
        for paragraph in doc.paragraphs:
            markdown_paragraph = self._convert_paragraph_to_markdown(paragraph, image_relationships)
            if markdown_paragraph.strip():
                markdown_parts.append(markdown_paragraph)
        
        # 处理表格
        for table in doc.tables:
            markdown_table = self._convert_table_to_markdown(table)
            if markdown_table.strip():
                markdown_parts.append(markdown_table)
        
        return markdown_parts
    
    def _create_legacy_doc_result(self, text_content: str, file_name: str, file_content: bytes) -> Dict[str, Any]:
        """创建旧格式doc文件的解析结果"""
        return {
            "text_content": text_content,
            "file_type": "legacy_word",
            "file_name": file_name,
            "file_size": len(file_content),
            "analysis_method": "legacy_doc_extraction",
            "content_analysis": self.analyze_text_comprehensive(text_content),
            "extraction_summary": {
                "success": True,
                "method": "legacy_doc_parser",
                "total_text_length": len(text_content)
            }
        } 