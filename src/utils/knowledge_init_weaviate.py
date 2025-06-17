#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
知识库初始化脚本 - Weaviate 版本
将 D:\knowledge_base 下的知识库内容转换为向量并存储到 Weaviate 中
根据 knowledge_init_weaviate.md 要求实现，使用 LangChain 框架进行 RAG
"""

import os
import re
import uuid
import json
import logging
import hashlib
import shutil
import base64
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from io import BytesIO

# 文件处理相关
import docx
from openpyxl import load_workbook
from PIL import Image
import pytesseract
from transformers import BlipProcessor, BlipForConditionalGeneration

# 向量化和嵌入
from sentence_transformers import SentenceTransformer

# Weaviate 和 LangChain
import weaviate
import weaviate.classes as wvc
from weaviate.auth import AuthApiKey
from langchain_community.vectorstores import Weaviate as LangChainWeaviate
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# 项目配置和 Redis
try:
    from ..resource.config import get_config
    from .redis_util import get_redis_manager
    from .weaviate_helper import get_weaviate_client
except ImportError:
    import sys
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    sys.path.insert(0, project_root)
    from src.resource.config import get_config
    from src.utils.redis_util import get_redis_manager
    from src.utils.weaviate_helper import get_weaviate_client

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class KnowledgeBaseInitializer:
    """知识库初始化器"""
    
    def __init__(self, knowledge_base_path: str = "D:\\knowledge_base"):
        """
        初始化知识库处理器
        
        Args:
            knowledge_base_path: 知识库根目录路径
        """
        # 初始化LLM客户端以使用_call_llm方法
        self.llm_client = None
        self._initialize_llm_client()
        
        self.knowledge_base_path = Path(knowledge_base_path)
        self.weaviate_client = None
        self.redis_manager = None
        self.embeddings_model = None
        self.langchain_embeddings = None
        self.langchain_vectorstore = None
        self.text_splitter = None
        self.blip_processor = None
        self.blip_model = None
        
        # 支持的文件类型
        self.supported_extensions = {'.docx', '.xlsx', '.java', '.xml'}
        
        # 图片输出目录 - 按要求设置为指定路径
        self.image_output_dir = self.knowledge_base_path / "链数_LS" / "需求文档" / "需求文档图片"
        
        # 初始化组件
        self._initialize_components()
    
    def _initialize_llm_client(self):
        """初始化LLM客户端"""
        try:
            from ..utils.volcengine_client import VolcengineClient
            config = get_config()
            volcengine_config = config.get_volcengine_config()
            self.llm_client = VolcengineClient(volcengine_config)
            logger.info("✅ LLM客户端初始化成功")
        except Exception as e:
            logger.warning(f"⚠️ LLM客户端初始化失败: {e}")
            self.llm_client = None
    
    def _call_llm(self, prompt: str, system_prompt: str = None, max_tokens: int = 4000) -> Optional[str]:
        """
        调用LLM进行分析
        
        Args:
            prompt: 用户提示
            system_prompt: 系统提示
            max_tokens: 最大token数
            
        Returns:
            模型响应文本
        """
        if not self.llm_client:
            logger.warning("LLM客户端未初始化")
            return None
        
        try:
            response = self.llm_client.chat(
                messages=[
                    {"role": "system", "content": system_prompt or "你是一个专业的文档分析助手"},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens
            )
            return response
        except Exception as e:
            logger.error(f"LLM调用失败: {e}")
            return None
    
    def _initialize_components(self):
        """初始化各种组件"""
        try:
            # 初始化 Weaviate 客户端
            self.weaviate_client = get_weaviate_client()
            logger.info("✅ Weaviate 客户端初始化成功")
            
            # 初始化 Redis 管理器
            self.redis_manager = get_redis_manager()
            logger.info("✅ Redis 管理器初始化成功")
            
            # 初始化嵌入模型 - 使用bge-large-zh（1024维，中文优化）
            self.embeddings_model = SentenceTransformer('BAAI/bge-large-zh')
            logger.info("✅ 嵌入模型加载成功 (bge-large-zh, 1024维，中文优化)")
            
            # 初始化 LangChain 嵌入模型
            self.langchain_embeddings = SentenceTransformerEmbeddings(
                model_name='BAAI/bge-large-zh'
            )
            logger.info("✅ LangChain 嵌入模型初始化成功")
            
            # 初始化文本分割器
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50,
                length_function=len,
                separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""]
            )
            logger.info("✅ 文本分割器初始化成功")
            
            # 创建图片输出目录
            self.image_output_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"✅ 图片输出目录创建成功: {self.image_output_dir}")
            
            logger.info("✅ 组件初始化完成")
            
        except Exception as e:
            logger.error(f"❌ 组件初始化失败: {e}")
            raise
    
    def _load_image_caption_model(self):
        """延迟加载图片描述模型"""
        if self.blip_processor is None:
            try:
                self.blip_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
                self.blip_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
                logger.info("✅ 图片描述模型加载成功")
            except Exception as e:
                logger.error(f"❌ 图片描述模型加载失败: {e}")
                self.blip_processor = None
                self.blip_model = None
    
    def _get_cache_key(self, file_path: str, cache_type: str) -> str:
        """生成缓存键 - 按要求使用指定格式"""
        return f"file:{file_path}:{cache_type}"
    
    def _get_cached_result(self, file_path: str, cache_type: str) -> Optional[Any]:
        """获取缓存结果"""
        try:
            cache_key = self._get_cache_key(file_path, cache_type)
            return self.redis_manager.get(cache_key, use_prefix=False)
        except Exception as e:
            logger.warning(f"获取缓存失败 {file_path}:{cache_type} - {e}")
            return None
    
    def _set_cached_result(self, file_path: str, cache_type: str, result: Any, ttl: int = 86400):
        """设置缓存结果"""
        try:
            cache_key = self._get_cache_key(file_path, cache_type)
            self.redis_manager.set(cache_key, result, ttl=ttl, use_prefix=False)
        except Exception as e:
            logger.warning(f"设置缓存失败 {file_path}:{cache_type} - {e}")
    
    def _extract_project_name(self, file_path: Path) -> str:
        """从文件路径提取项目名称"""
        try:
            # 获取相对于知识库根目录的路径
            relative_path = file_path.relative_to(self.knowledge_base_path)
            parts = relative_path.parts
            
            # 按要求，将目录结构存储为LS
            if len(parts) >= 1:
                # 检查是否包含特定的项目标识
                for part in parts:
                    if 'zqyl-ls' in part.lower():
                        return 'zqyl-ls'
                    elif 'ls' in part.lower() or '链数' in part:
                        return 'LS'
            
            # 默认返回LS
            return "LS"
            
        except Exception as e:
            logger.warning(f"提取项目名称失败 {file_path}: {e}")
            return "LS"
    
    def _clear_weaviate_database(self):
        """清空Weaviate数据库所有数据"""
        try:
            logger.info("🗑️ 开始清空Weaviate数据库...")
            
            # 获取所有集合
            collections = self.weaviate_client.collections.list_all()
            
            for collection_name in collections:
                try:
                    # 删除集合
                    self.weaviate_client.collections.delete(collection_name)
                    logger.info(f"✅ 删除集合: {collection_name}")
                except Exception as e:
                    logger.warning(f"⚠️ 删除集合失败 {collection_name}: {e}")
            
            logger.info("✅ Weaviate数据库清空完成")
            
        except Exception as e:
            logger.error(f"❌ 清空Weaviate数据库失败: {e}")
            raise
    
    def _docx_to_markdown(self, doc, file_path: Path) -> Tuple[str, List[Dict[str, Any]]]:
        """将DOCX转换为Markdown格式，并提取图片信息"""
        markdown_content = []
        image_info = []
        
        # 确保图片输出目录存在
        self.image_output_dir.mkdir(parents=True, exist_ok=True)
        
        try:
            # 首先提取所有图片
            image_info = self._extract_images_from_docx(doc, file_path)
            
            # 处理段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检测标题
                if self._is_heading(paragraph, text):
                    # 根据标题级别添加markdown标记
                    level = self._get_heading_level(paragraph, text)
                    markdown_content.append(f"{'#' * level} {text}")
                else:
                    # 普通段落
                    markdown_content.append(text)
            
            # 如果提取到图片，在内容末尾添加图片引用
            if image_info:
                markdown_content.append("\n## 文档图片")
                for img in image_info:
                    markdown_content.append(f"![{img['name']}]({img['path']})")
            
            return "\n\n".join(markdown_content), image_info
            
        except Exception as e:
            logger.error(f"DOCX转Markdown失败 {file_path}: {e}")
            return "", []
    
    def _extract_images_from_docx(self, doc, file_path: Path) -> List[Dict[str, Any]]:
        """从DOCX文档中提取图片"""
        image_info = []
        image_counter = 1
        
        try:
            # 获取文档的关系部分
            document_part = doc.part
            
            # 遍历所有关系，查找图片
            for rel_id, relationship in document_part.rels.items():
                if "image" in relationship.target_ref:
                    try:
                        # 获取图片数据
                        image_part = relationship.target_part
                        image_data = image_part.blob
                        
                        # 确定图片扩展名
                        content_type = image_part.content_type
                        if 'png' in content_type:
                            ext = '.png'
                        elif 'jpeg' in content_type or 'jpg' in content_type:
                            ext = '.jpg'
                        elif 'gif' in content_type:
                            ext = '.gif'
                        else:
                            ext = '.png'  # 默认使用png
                        
                        # 生成图片文件名和路径
                        image_name = f"{file_path.stem}_image_{image_counter:03d}{ext}"
                        image_path = self.image_output_dir / image_name
                        
                        # 保存并可能压缩图片
                        original_size = len(image_data)
                        if original_size > 1024 * 1024:  # 大于1MB的图片进行压缩
                            try:
                                # 先保存原图
                                with open(image_path, 'wb') as img_file:
                                    img_file.write(image_data)
                                
                                # 使用PIL压缩图片
                                from PIL import Image
                                with Image.open(image_path) as img:
                                    # 如果图片很大，调整尺寸
                                    if img.width > 1920 or img.height > 1080:
                                        img.thumbnail((1920, 1080), Image.Resampling.LANCZOS)
                                    
                                    # 保存压缩后的图片
                                    if ext.lower() in ['.jpg', '.jpeg']:
                                        img.save(image_path, 'JPEG', quality=85, optimize=True)
                                    else:
                                        img.save(image_path, 'PNG', optimize=True)
                                
                                compressed_size = image_path.stat().st_size
                                compression_ratio = (1 - compressed_size / original_size) * 100
                                logger.info(f"✅ 提取并压缩图片: {image_name} (压缩率: {compression_ratio:.1f}%)")
                            except Exception as compress_error:
                                # 压缩失败，保存原图
                                with open(image_path, 'wb') as img_file:
                                    img_file.write(image_data)
                                logger.warning(f"⚠️ 图片压缩失败，保存原图: {image_name} - {compress_error}")
                        else:
                            # 小图片直接保存
                            with open(image_path, 'wb') as img_file:
                                img_file.write(image_data)
                            logger.info(f"✅ 提取图片: {image_name}")
                        
                        image_info.append({
                            'name': image_name,
                            'path': str(image_path),
                            'rel_id': rel_id,
                            'content_type': content_type,
                            'original_size': original_size,
                            'final_size': image_path.stat().st_size
                        })
                        
                        image_counter += 1
                        
                    except Exception as e:
                        logger.warning(f"⚠️ 提取图片失败 (rel_id: {rel_id}): {e}")
            
            if image_info:
                logger.info(f"📸 从 {file_path.name} 提取了 {len(image_info)} 张图片")
            
        except Exception as e:
            logger.error(f"❌ 图片提取过程失败 {file_path}: {e}")
        
        return image_info
    
    def _extract_text_from_damaged_docx(self, file_path: Path) -> str:
        """从损坏的DOCX文件中提取文本内容"""
        import zipfile
        import xml.etree.ElementTree as ET
        
        try:
            text_content = []
            
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # 列出所有文件
                file_list = docx_zip.namelist()
                logger.info(f"📋 DOCX文件内容: {file_list}")
                
                # 尝试读取主文档内容
                main_doc_files = [
                    'word/document.xml',
                    'word/document2.xml', 
                    'document.xml'
                ]
                
                for doc_file in main_doc_files:
                    if doc_file in file_list:
                        try:
                            with docx_zip.open(doc_file) as xml_file:
                                xml_content = xml_file.read()
                                
                                # 解析XML并提取文本
                                root = ET.fromstring(xml_content)
                                
                                # 定义命名空间
                                namespaces = {
                                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                }
                                
                                # 提取所有文本节点
                                for text_elem in root.findall('.//w:t', namespaces):
                                    if text_elem.text:
                                        text_content.append(text_elem.text)
                                
                                logger.info(f"✅ 成功从 {doc_file} 提取文本")
                                break
                                
                        except Exception as e:
                            logger.warning(f"⚠️ 无法处理 {doc_file}: {e}")
                            continue
                
                # 如果主文档失败，尝试其他可能的文本文件
                if not text_content:
                    for file_name in file_list:
                        if file_name.endswith('.xml') and 'word' in file_name:
                            try:
                                with docx_zip.open(file_name) as xml_file:
                                    xml_content = xml_file.read().decode('utf-8', errors='ignore')
                                    # 简单的正则提取文本
                                    import re
                                    text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                                    if text_matches:
                                        text_content.extend(text_matches)
                                        logger.info(f"✅ 通过正则从 {file_name} 提取文本")
                                        break
                            except Exception as e:
                                continue
            
            if text_content:
                # 清理和组织文本
                cleaned_text = []
                current_line = ""
                
                for text in text_content:
                    text = text.strip()
                    if not text:
                        continue
                    
                    # 检查是否是新段落的开始
                    if any(char in text for char in ['.', '。', '!', '！', '?', '？']) and len(current_line) > 20:
                        current_line += text
                        cleaned_text.append(current_line)
                        current_line = ""
                    else:
                        current_line += text + " "
                
                # 添加最后一行
                if current_line.strip():
                    cleaned_text.append(current_line.strip())
                
                final_text = f"# {file_path.name}\n\n" + "\n\n".join(cleaned_text)
                
                logger.info(f"📝 成功恢复文本内容，共 {len(cleaned_text)} 段落")
                return final_text
            
        except Exception as e:
            logger.error(f"❌ zipfile方法也失败: {e}")
        
        return ""
    
    def _process_docx_with_pandoc(self, file_path: Path) -> List[Document]:
        """使用Pandoc处理DOCX文件"""
        try:
            import pypandoc
            
            # 确保Pandoc可用
            try:
                pypandoc.get_pandoc_version()
            except OSError:
                logger.info("🔄 Pandoc未找到，尝试自动下载...")
                try:
                    pypandoc.download_pandoc()
                    logger.info("✅ Pandoc下载完成")
                except Exception as download_error:
                    logger.error(f"❌ Pandoc下载失败: {download_error}")
                    return []
            
            # 确保图片输出目录存在
            self.image_output_dir.mkdir(parents=True, exist_ok=True)
            
            # 使用Pandoc将DOCX转换为Markdown
            logger.info(f"📄 使用Pandoc转换DOCX到Markdown...")
            
            # 设置Pandoc参数
            extra_args = [
                '--extract-media', str(self.image_output_dir.parent),  # 提取图片到指定目录
                '--wrap=none',  # 不自动换行
                '--markdown-headings=atx'  # 使用ATX样式标题
            ]
            
            try:
                # 尝试转换为markdown
                markdown_content = pypandoc.convert_file(
                    str(file_path), 
                    'markdown',
                    extra_args=extra_args
                )
                
                if not markdown_content.strip():
                    logger.warning(f"⚠️ Pandoc转换结果为空")
                    return []
                
                logger.info(f"✅ Pandoc转换成功，内容长度: {len(markdown_content)}")
                
                # 处理提取的图片
                image_info = self._process_extracted_images(file_path)
                
                # 按章节分割内容
                sections = self._split_pandoc_markdown(markdown_content, file_path)
                
                documents = []
                for i, section in enumerate(sections):
                    doc_obj = Document(
                        page_content=section['content'],
                        metadata={
                            'file_path': str(file_path),
                            'file_name': file_path.name,
                            'project': self._extract_project_name(file_path),
                            'file_type': 'docx',
                            'source_type': 'pandoc_markdown',
                            'section': section['title'],
                            'chunk_index': i,
                            'total_chunks': len(sections),
                            'processor': 'pandoc'
                        }
                    )
                    documents.append(doc_obj)
                
                # 添加图片文档
                for img_info in image_info:
                    img_doc = self._process_image_file(Path(img_info['path']), file_path)
                    if img_doc:
                        documents.append(img_doc)
                
                # 缓存结果
                cache_data = [{'page_content': doc.page_content, 'metadata': doc.metadata} for doc in documents]
                self._set_cached_result(str(file_path), "docx_content", cache_data)
                
                logger.info(f"✅ Pandoc处理完成: {file_path.name} ({len(documents)} 个文档)")
                return documents
                
            except Exception as convert_error:
                logger.error(f"❌ Pandoc转换失败: {convert_error}")
                
                # 尝试简单文本提取
                try:
                    text_content = pypandoc.convert_file(str(file_path), 'plain')
                    if text_content.strip():
                        simple_doc = Document(
                            page_content=f"# {file_path.name}\n\n{text_content}",
                            metadata={
                                'file_path': str(file_path),
                                'file_name': file_path.name,
                                'project': self._extract_project_name(file_path),
                                'file_type': 'docx',
                                'source_type': 'pandoc_plain',
                                'section': '文档内容',
                                'processor': 'pandoc_fallback'
                            }
                        )
                        logger.info(f"✅ Pandoc简单文本提取成功: {file_path.name}")
                        return [simple_doc]
                except Exception as plain_error:
                    logger.error(f"❌ Pandoc简单文本提取也失败: {plain_error}")
                
                return []
                
        except ImportError:
            logger.error("❌ pypandoc未安装，无法使用Pandoc处理")
            return []
        except Exception as e:
            logger.error(f"❌ Pandoc处理异常: {e}")
            return []
    
    def _process_extracted_images(self, file_path: Path) -> List[Dict[str, Any]]:
        """处理Pandoc提取的图片"""
        image_info = []
        
        try:
            # Pandoc会将图片提取到media目录
            media_dir = self.image_output_dir.parent / 'media'
            if media_dir.exists():
                image_counter = 1
                for img_file in media_dir.glob('**/*'):
                    if img_file.is_file() and img_file.suffix.lower() in ['.png', '.jpg', '.jpeg', '.gif']:
                        # 移动图片到我们的图片目录并重命名
                        new_name = f"{file_path.stem}_pandoc_image_{image_counter:03d}{img_file.suffix}"
                        new_path = self.image_output_dir / new_name
                        
                        try:
                            # 复制并可能压缩图片
                            import shutil
                            shutil.copy2(img_file, new_path)
                            
                            # 压缩大图片
                            if new_path.stat().st_size > 1024 * 1024:  # 大于1MB
                                self._compress_image(new_path)
                            
                            image_info.append({
                                'name': new_name,
                                'path': str(new_path),
                                'original_path': str(img_file),
                                'size': new_path.stat().st_size
                            })
                            
                            image_counter += 1
                            logger.info(f"📸 处理Pandoc提取的图片: {new_name}")
                            
                        except Exception as e:
                            logger.warning(f"⚠️ 处理图片失败 {img_file}: {e}")
                
                # 清理media目录
                try:
                    shutil.rmtree(media_dir)
                except:
                    pass
                    
        except Exception as e:
            logger.warning(f"⚠️ 处理提取图片时出错: {e}")
        
        return image_info
    
    def _compress_image(self, image_path: Path):
        """压缩图片"""
        try:
            from PIL import Image
            with Image.open(image_path) as img:
                # 限制最大尺寸
                if img.width > 1920 or img.height > 1080:
                    img.thumbnail((1920, 1080), Image.Resampling.LANCZOS)
                
                # 保存压缩后的图片
                if image_path.suffix.lower() in ['.jpg', '.jpeg']:
                    img.save(image_path, 'JPEG', quality=85, optimize=True)
                else:
                    img.save(image_path, 'PNG', optimize=True)
                    
        except Exception as e:
            logger.warning(f"⚠️ 图片压缩失败: {e}")
    
    def _split_pandoc_markdown(self, markdown_content: str, file_path: Path) -> List[Dict[str, Any]]:
        """分割Pandoc生成的Markdown内容"""
        sections = []
        lines = markdown_content.split('\n')
        
        current_title = "文档开始"
        current_content = []
        
        for line in lines:
            line = line.strip()
            
            # 检测标题行
            if line.startswith('#'):
                # 保存之前的章节
                if current_content:
                    sections.append({
                        'title': current_title,
                        'content': '\n'.join(current_content).strip()
                    })
                
                # 开始新章节
                current_title = line.lstrip('#').strip() or "无标题章节"
                current_content = []
            else:
                if line:  # 非空行
                    current_content.append(line)
        
        # 处理最后一个章节
        if current_content:
            sections.append({
                'title': current_title,
                'content': '\n'.join(current_content).strip()
            })
        
        # 如果没有找到章节，创建一个默认章节
        if not sections:
            sections.append({
                'title': file_path.stem,
                'content': markdown_content
            })
        
        return sections
    
    def _is_heading(self, paragraph, text: str) -> bool:
        """判断是否为标题"""
        # 检查段落样式
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # 检查文本格式（数字开头的标题）
        if re.match(r'^\d+\.?\d*\s+', text):
            return True
        
        # 检查是否为短文本且可能是标题
        if len(text) < 50 and any(keyword in text for keyword in ['系统', '模块', '功能', '接口', '设计', '架构']):
            return True
        
        return False
    
    def _get_heading_level(self, paragraph, text: str) -> int:
        """获取标题级别"""
        if paragraph.style.name.startswith('Heading'):
            try:
                return int(paragraph.style.name.replace('Heading ', ''))
            except:
                return 1
        
        # 根据数字格式判断级别
        match = re.match(r'^(\d+)\.?(\d*)\s+', text)
        if match:
            if match.group(2):  # 有二级数字，如 3.1
                return 2
            else:  # 只有一级数字，如 3
                return 1
        
        return 1
    
    def _split_markdown_by_sections(self, markdown_content: str, file_path: Path) -> List[Dict[str, Any]]:
        """按标题分割Markdown内容"""
        sections = []
        lines = markdown_content.split('\n')
        
        current_section = ""
        current_content = []
        current_images = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测标题行
            if line.startswith('#'):
                # 保存之前的段落
                if current_section and current_content:
                    sections.append({
                        'section': current_section,
                        'content': '\n'.join(current_content),
                        'image_refs': current_images.copy()
                    })
                
                # 开始新段落
                current_section = line.lstrip('#').strip()
                current_content = []
                current_images = []
            
            elif line.startswith('!['):
                # 图片引用
                image_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
                if image_match:
                    current_images.append(image_match.group(2))
            
            else:
                # 普通内容
                current_content.append(line)
        
        # 处理最后一个段落
        if current_section and current_content:
            sections.append({
                'section': current_section,
                'content': '\n'.join(current_content),
                'image_refs': current_images
            })
        
        return sections
    
    def _process_docx_file(self, file_path: Path) -> List[Document]:
        """处理 DOCX 文件 - 返回LangChain Document对象"""
        documents = []
        
        try:
            # 检查缓存
            cached_result = self._get_cached_result(str(file_path), "docx_content")
            if cached_result:
                logger.info(f"📋 使用缓存的 DOCX 内容: {file_path.name}")
                return [Document(**doc) for doc in cached_result]
            
            # 尝试使用多种方法处理DOCX文件
            doc = None
            markdown_content = ""
            image_info = []
            
            # 方法1: 尝试使用Pandoc处理（推荐，容错性最好）
            try:
                logger.info(f"🔄 尝试使用Pandoc处理DOCX文件: {file_path.name}")
                pandoc_result = self._process_docx_with_pandoc(file_path)
                if pandoc_result:
                    return pandoc_result
            except Exception as pandoc_error:
                logger.warning(f"⚠️ Pandoc处理失败: {pandoc_error}")
            
                # 方法2: 尝试使用python-docx
                try:
                    doc = docx.Document(file_path)
                    logger.info(f"✅ python-docx成功打开文件: {file_path.name}")
                except Exception as docx_error:
                    logger.warning(f"⚠️ python-docx无法打开文件 {file_path}: {docx_error}")
                    
                    # 方法3: 尝试使用zipfile直接提取文本内容
                    try:
                        logger.info(f"🔄 尝试使用zipfile方法处理损坏的DOCX文件...")
                        text_content = self._extract_text_from_damaged_docx(file_path)
                        if text_content:
                            # 创建基于提取文本的文档
                            recovered_doc = Document(
                                page_content=text_content,
                                metadata={
                                    'file_path': str(file_path),
                                    'file_name': file_path.name,
                                    'project': self._extract_project_name(file_path),
                                    'file_type': 'docx',
                                    'source_type': 'recovered_text',
                                    'section': '恢复的文档内容',
                                    'status': 'recovered'
                                }
                            )
                            
                            # 缓存结果
                            cache_data = [{'page_content': recovered_doc.page_content, 'metadata': recovered_doc.metadata}]
                            self._set_cached_result(str(file_path), "docx_content", cache_data)
                            
                            logger.info(f"✅ 从损坏的DOCX文件恢复文本内容: {file_path.name}")
                            return [recovered_doc]
                        
                    except Exception as recovery_error:
                        logger.error(f"❌ 文本恢复也失败: {recovery_error}")
                    
                # 所有方法都失败，创建错误信息文档
                logger.error(f"❌ 完全无法处理DOCX文件 {file_path}")
                error_doc = Document(
                    page_content=f"文档: {file_path.name}\n状态: 文件损坏，无法读取\n错误: {str(docx_error)}\n路径: {file_path}\n\n建议: 请使用Word重新保存此文件",
                    metadata={
                        'file_path': str(file_path),
                        'file_name': file_path.name,
                        'project': self._extract_project_name(file_path),
                        'file_type': 'docx',
                        'source_type': 'error',
                        'section': '文档信息',
                        'status': 'corrupted'
                    }
                )
                return [error_doc]
            
            # 将DOCX转换为Markdown
            markdown_content, image_info = self._docx_to_markdown(doc, file_path)
            
            # 按标题分割内容
            sections = self._split_markdown_by_sections(markdown_content, file_path)
            
            for section in sections:
                # 创建LangChain Document对象
                document = Document(
                    page_content=section['content'],
                    metadata={
                        'file_path': str(file_path),
                        'file_name': file_path.name,
                        'project': self._extract_project_name(file_path),
                        'file_type': 'docx',
                        'source_type': 'text',
                        'section': section['section'],
                        'image_refs': section.get('image_refs', [])
                    }
                )
                documents.append(document)
            
            # 处理图片
            for img_info in image_info:
                img_doc = self._process_image_file(Path(img_info['path']), file_path)
                if img_doc:
                    documents.append(img_doc)
            
            # 缓存结果
            cache_data = [{'page_content': doc.page_content, 'metadata': doc.metadata} for doc in documents]
            self._set_cached_result(str(file_path), "docx_content", cache_data)
            
            logger.info(f"✅ DOCX 文件处理完成: {file_path.name} ({len(documents)} 个文档)")
            
        except Exception as e:
            logger.error(f"❌ DOCX 文件处理失败 {file_path}: {e}")
        
        return documents
    
    def _process_image_file(self, image_path: Path, source_file: Path) -> Optional[Document]:
        """处理图片文件 - OCR和描述生成"""
        try:
            if not image_path.exists():
                return None
            
            # 检查缓存
            cached_result = self._get_cached_result(str(image_path), "image_analysis")
            if cached_result:
                return Document(**cached_result)
            
            # 延迟加载图片描述模型
            self._load_image_caption_model()
            
            # 加载图片
            image = Image.open(image_path)
            
            # OCR文本提取
            ocr_text = ""
            try:
                ocr_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            except Exception as e:
                logger.warning(f"OCR失败 {image_path}: {e}")
            
            # 图片描述生成
            caption = ""
            if self.blip_processor and self.blip_model:
                try:
                    inputs = self.blip_processor(image, return_tensors="pt")
                    out = self.blip_model.generate(**inputs, max_length=50)
                    caption = self.blip_processor.decode(out[0], skip_special_tokens=True)
                except Exception as e:
                    logger.warning(f"图片描述生成失败 {image_path}: {e}")
            
            # 合并文本内容
            combined_text = f"图片描述: {caption}\nOCR文本: {ocr_text}".strip()
            
            # 创建Document对象
            document = Document(
                page_content=combined_text,
                metadata={
                    'file_path': str(source_file),
                    'file_name': source_file.name,
                    'project': self._extract_project_name(source_file),
                    'file_type': 'image',
                    'source_type': 'image_description',
                    'image_path': str(image_path),
                    'ocr_text': ocr_text,
                    'caption': caption
                }
            )
            
            # 缓存结果
            cache_data = {'page_content': document.page_content, 'metadata': document.metadata}
            self._set_cached_result(str(image_path), "image_analysis", cache_data)
            
            return document
            
        except Exception as e:
            logger.error(f"图片处理失败 {image_path}: {e}")
            return None
    
    def _process_xlsx_file(self, file_path: Path) -> List[Document]:
        """处理 XLSX 文件"""
        documents = []
        
        try:
            # 检查缓存
            cached_result = self._get_cached_result(str(file_path), "xlsx_content")
            if cached_result:
                logger.info(f"📋 使用缓存的 XLSX 内容: {file_path.name}")
                return [Document(**doc) for doc in cached_result]
            
            workbook = load_workbook(file_path, read_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # 提取表格内容
                sheet_content = []
                headers = []
                
                # 获取表头
                for cell in sheet[1]:
                    if cell.value:
                        headers.append(str(cell.value))
                
                if headers:
                    sheet_content.append("表头: " + " | ".join(headers))
                
                # 获取数据行
                for row in sheet.iter_rows(min_row=2, values_only=True):
                    row_data = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_data.append(str(cell_value))
                    
                    if row_data:
                        sheet_content.append(" | ".join(row_data))
                
                if sheet_content:
                    full_content = f"工作表: {sheet_name}\n" + "\n".join(sheet_content)
                    
                    # 分割内容
                    text_chunks = self.text_splitter.split_text(full_content)
                    
                    for i, chunk in enumerate(text_chunks):
                        document = Document(
                            page_content=chunk,
                            metadata={
                                'file_path': str(file_path),
                                'file_name': file_path.name,
                                'project': self._extract_project_name(file_path),
                                'file_type': 'xlsx',
                                'source_type': 'excel',
                                'sheet_name': sheet_name,
                                'chunk_index': i
                            }
                        )
                        documents.append(document)
            
            workbook.close()
            
            # 缓存结果
            cache_data = [{'page_content': doc.page_content, 'metadata': doc.metadata} for doc in documents]
            self._set_cached_result(str(file_path), "xlsx_content", cache_data)
            
            logger.info(f"✅ XLSX 文件处理完成: {file_path.name} ({len(documents)} 个文档)")
            
        except Exception as e:
            logger.error(f"❌ XLSX 文件处理失败 {file_path}: {e}")
        
        return documents
    
    def _process_java_file(self, file_path: Path) -> List[Document]:
        """处理 Java 文件"""
        documents = []
        
        try:
            # 检查缓存
            cached_result = self._get_cached_result(str(file_path), "java_content")
            if cached_result:
                logger.info(f"📋 使用缓存的 Java 内容: {file_path.name}")
                return [Document(**doc) for doc in cached_result]
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # 按类和方法分割Java代码
            code_blocks = self._split_java_code(content)
            
            for i, block in enumerate(code_blocks):
                if len(block.strip()) > 50:  # 过滤太短的内容
                    document = Document(
                        page_content=block.strip(),
                        metadata={
                            'file_path': str(file_path),
                            'file_name': file_path.name,
                            'project': self._extract_project_name(file_path),
                            'file_type': 'java',
                            'source_type': 'code',
                            'chunk_index': i
                        }
                    )
                    documents.append(document)
            
            # 缓存结果
            cache_data = [{'page_content': doc.page_content, 'metadata': doc.metadata} for doc in documents]
            self._set_cached_result(str(file_path), "java_content", cache_data)
            
            logger.info(f"✅ Java 文件处理完成: {file_path.name} ({len(documents)} 个文档)")
            
        except Exception as e:
            logger.error(f"❌ Java 文件处理失败 {file_path}: {e}")
        
        return documents
    
    def _split_java_code(self, content: str) -> List[str]:
        """分割Java代码"""
        blocks = []
        
        # 按类分割
        class_pattern = r'(public\s+class\s+\w+.*?(?=public\s+class|\Z))'
        classes = re.findall(class_pattern, content, re.DOTALL)
        
        if classes:
            blocks.extend(classes)
        else:
            # 如果没有找到类，按方法分割
            method_pattern = r'(public\s+\w+.*?\{.*?\n\s*\})'
            methods = re.findall(method_pattern, content, re.DOTALL)
            
            if methods:
                blocks.extend(methods)
            else:
                # 如果都没有找到，按文本分割
                text_chunks = self.text_splitter.split_text(content)
                blocks.extend(text_chunks)
        
        return blocks
    
    def _process_xml_file(self, file_path: Path) -> List[Document]:
        """处理 XML 文件"""
        documents = []
        
        try:
            # 检查缓存
            cached_result = self._get_cached_result(str(file_path), "xml_content")
            if cached_result:
                logger.info(f"📋 使用缓存的 XML 内容: {file_path.name}")
                return [Document(**doc) for doc in cached_result]
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # 按文本分割 XML 内容
            text_chunks = self.text_splitter.split_text(content)
            
            for i, chunk in enumerate(text_chunks):
                document = Document(
                    page_content=chunk,
                    metadata={
                        'file_path': str(file_path),
                        'file_name': file_path.name,
                        'project': self._extract_project_name(file_path),
                        'file_type': 'xml',
                        'source_type': 'config',
                        'chunk_index': i
                    }
                )
                documents.append(document)
            
            # 缓存结果
            cache_data = [{'page_content': doc.page_content, 'metadata': doc.metadata} for doc in documents]
            self._set_cached_result(str(file_path), "xml_content", cache_data)
            
            logger.info(f"✅ XML 文件处理完成: {file_path.name} ({len(documents)} 个文档)")
            
        except Exception as e:
            logger.error(f"❌ XML 文件处理失败 {file_path}: {e}")
        
        return documents
    
    def _create_weaviate_schema(self):
        """创建 Weaviate 模式 - 按要求定义Document类"""
        try:
            collection_name = "Document"
            
            # 检查集合是否已存在
            if self.weaviate_client.collections.exists(collection_name):
                logger.info(f"📋 集合 {collection_name} 已存在，跳过创建")
                return collection_name
            
            # 创建集合 - 按要求设置vectorizer: none
            self.weaviate_client.collections.create(
                name=collection_name,
                vectorizer_config=wvc.config.Configure.Vectorizer.none(),
                properties=[
                    wvc.config.Property(name="content", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="file_path", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="file_name", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="project", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="file_type", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="source_type", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="image_path", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="section", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="sheet_name", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="chunk_index", data_type=wvc.config.DataType.INT),
                    wvc.config.Property(name="created_at", data_type=wvc.config.DataType.DATE),
                    wvc.config.Property(name="ocr_text", data_type=wvc.config.DataType.TEXT),
                    wvc.config.Property(name="caption", data_type=wvc.config.DataType.TEXT)
                ]
            )
            
            logger.info(f"✅ Weaviate 集合创建成功: {collection_name}")
            return collection_name
            
        except Exception as e:
            logger.error(f"❌ Weaviate 集合创建失败: {e}")
            raise
    
    def _initialize_langchain_vectorstore(self, collection_name: str):
        """初始化LangChain向量存储 - 使用手动RAG实现"""
        try:
            # 由于版本兼容性问题，使用手动RAG实现
            logger.info("✅ 使用手动RAG实现 (跳过LangChain向量存储)")
            self.langchain_vectorstore = None
            
        except Exception as e:
            logger.error(f"❌ LangChain 向量存储初始化失败: {e}")
            self.langchain_vectorstore = None
    
    def _insert_documents_batch(self, documents: List[Document], collection_name: str):
        """手动批量插入文档到Weaviate"""
        try:
            collection = self.weaviate_client.collections.get(collection_name)
            
            # 准备批量插入的数据
            objects = []
            for doc in documents:
                # 生成向量
                vector = self.embeddings_model.encode(doc.page_content).tolist()
                
                # 准备对象数据
                obj_data = {
                    "content": doc.page_content,
                    "vector": vector,
                    **doc.metadata  # 展开所有元数据
                }
                objects.append(obj_data)
            
            # 批量插入
            with collection.batch.dynamic() as batch:
                for obj in objects:
                    batch.add_object(
                        properties={k: v for k, v in obj.items() if k != "vector"},
                        vector=obj["vector"]
                    )
            
        except Exception as e:
            logger.error(f"❌ 批量插入文档失败: {e}")
            raise
    
    def scan_knowledge_base(self) -> List[Path]:
        """扫描知识库目录，获取所有支持的文件"""
        files = []
        
        try:
            for file_path in self.knowledge_base_path.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                    files.append(file_path)
            
            logger.info(f"📁 扫描完成，找到 {len(files)} 个支持的文件")
            
            # 按文件类型统计
            stats = {}
            for file_path in files:
                ext = file_path.suffix.lower()
                stats[ext] = stats.get(ext, 0) + 1
            
            for ext, count in stats.items():
                logger.info(f"   {ext}: {count} 个文件")
            
        except Exception as e:
            logger.error(f"❌ 扫描知识库失败: {e}")
        
        return files
    
    def process_files(self, files: List[Path]) -> List[Document]:
        """处理所有文件"""
        all_documents = []
        
        for i, file_path in enumerate(files, 1):
            logger.info(f"🔄 处理文件 {i}/{len(files)}: {file_path.name}")
            
            try:
                documents = []
                
                if file_path.suffix.lower() == '.docx':
                    documents = self._process_docx_file(file_path)
                elif file_path.suffix.lower() == '.xlsx':
                    documents = self._process_xlsx_file(file_path)
                elif file_path.suffix.lower() == '.java':
                    documents = self._process_java_file(file_path)
                elif file_path.suffix.lower() == '.xml':
                    documents = self._process_xml_file(file_path)
                
                all_documents.extend(documents)
                
            except Exception as e:
                logger.error(f"❌ 处理文件失败 {file_path}: {e}")
        
        logger.info(f"📊 文件处理完成，总计生成 {len(all_documents)} 个文档")
        return all_documents
    
    def initialize_knowledge_base(self):
        """初始化知识库"""
        try:
            logger.info("🚀 开始初始化知识库...")
            
            # 1. 清空Weaviate数据库
            self._clear_weaviate_database()
            
            # 2. 检查知识库目录
            if not self.knowledge_base_path.exists():
                raise FileNotFoundError(f"知识库目录不存在: {self.knowledge_base_path}")
            
            # 3. 扫描文件
            files = self.scan_knowledge_base()
            if not files:
                logger.warning("⚠️ 没有找到支持的文件")
                return
            
            # 4. 创建 Weaviate 模式
            collection_name = self._create_weaviate_schema()
            
            # 5. 初始化LangChain向量存储
            self._initialize_langchain_vectorstore(collection_name)
            
            # 6. 处理文件
            documents = self.process_files(files)
            if not documents:
                logger.warning("⚠️ 没有生成任何文档")
                return
            
            # 7. 手动批量插入到Weaviate - 按要求batch_size=100
            logger.info("📤 开始批量插入文档到Weaviate...")
            batch_size = 100
            
            for i in range(0, len(documents), batch_size):
                batch = documents[i:i + batch_size]
                try:
                    # 手动插入文档
                    self._insert_documents_batch(batch, collection_name)
                    logger.info(f"✅ 批次 {i//batch_size + 1}/{(len(documents) + batch_size - 1)//batch_size} 插入成功 ({len(batch)} 个文档)")
                except Exception as e:
                    logger.error(f"❌ 批次 {i//batch_size + 1} 插入失败: {e}")
            
            logger.info("🎉 知识库初始化完成！")
            
        except Exception as e:
            logger.error(f"❌ 知识库初始化失败: {e}")
            raise
    
    def query_knowledge_base(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """手动查询知识库"""
        try:
            # 生成查询向量
            query_vector = self.embeddings_model.encode(query).tolist()
            
            # 使用Weaviate进行向量搜索
            collection = self.weaviate_client.collections.get("Document")
            
            results = collection.query.near_vector(
                near_vector=query_vector,
                limit=k,
                return_metadata=["distance"]
            )
            
            # 格式化结果
            formatted_results = []
            for obj in results.objects:
                formatted_results.append({
                    'content': obj.properties.get('content', ''),
                    'metadata': {k: v for k, v in obj.properties.items() if k != 'content'},
                    'similarity_score': 1 - obj.metadata.distance,  # 转换为相似度分数
                    'distance': obj.metadata.distance
                })
            
            return formatted_results
            
        except Exception as e:
            logger.error(f"❌ 查询知识库失败: {e}")
            return []
    
    def close(self):
        """关闭连接"""
        if self.weaviate_client:
            self.weaviate_client.close()
        if self.redis_manager:
            self.redis_manager.close()

def main():
    """主函数 - 检查特定文档"""
    initializer = None
    
    try:
        # 创建初始化器
        initializer = KnowledgeBaseInitializer()
        
        # 检查特定文档
        file_name = "LS-1(YS-72)_需求文档-链数一期V1.8.docx"
        print(f"🔍 检查文档: {file_name}")
       

        collection = initializer.weaviate_client.collections.get("Document")
        print(f"🔍 查询出当前文档: {collection}")

        # 使用正确的查询语法，使用filters参数而不是where
        existing_docs = collection.query.fetch_objects(
            filters=wvc.query.Filter.by_property("file_name").equal(file_name)
        )
        print(f"🔍 查询出当前文件名文档: {existing_docs}")
        
        # 先获取所有文档，然后过滤
        print("📋 正在查询数据库...")
        all_results = collection.query.fetch_objects(limit=1000)
        
        if not all_results.objects:
            print("❌ 数据库为空，没有任何向量化文档")
            return
        
        # 手动过滤特定文件
        target_docs = []
        all_files = set()
        
        for obj in all_results.objects:
            props = obj.properties
            current_file = props.get('file_name', 'unknown')
            all_files.add(current_file)
            
            if file_name in current_file or current_file == file_name:
                target_docs.append(obj)
        
        if not target_docs:
            print(f"❌ 未找到文件 '{file_name}' 的向量化记录")
            print(f"\n📋 数据库中的所有文件 (共{len(all_files)}个):")
            for fname in sorted(all_files):
                print(f"  - {fname}")
        else:
            print(f"✅ 找到文件 '{file_name}' 的 {len(target_docs)} 个向量化文档块")
            
            # 统计信息
            processors = set()
            source_types = set()
            sections = set()
            
            # 显示详细信息
            for i, obj in enumerate(target_docs[:5]):  # 只显示前5个
                props = obj.properties
                print(f"\n{i+1}. 章节: {props.get('section', 'N/A')}")
                print(f"   处理器: {props.get('processor', 'N/A')}")
                print(f"   来源类型: {props.get('source_type', 'N/A')}")
                print(f"   文件类型: {props.get('file_type', 'N/A')}")
                print(f"   项目: {props.get('project', 'N/A')}")
                print(f"   内容预览: {props.get('content', '')[:200]}...")
                
                processors.add(props.get('processor', 'unknown'))
                source_types.add(props.get('source_type', 'unknown'))
                if props.get('section'):
                    sections.add(props.get('section'))
            
            if len(target_docs) > 5:
                print(f"\n... 还有 {len(target_docs) - 5} 个文档块")
            
            print(f"\n📊 统计信息:")
            print(f"  - 总文档块数: {len(target_docs)}")
            print(f"  - 使用的处理器: {list(processors)}")
            print(f"  - 来源类型: {list(source_types)}")
            print(f"  - 章节数量: {len(sections)}")
            
            if sections:
                print(f"  - 章节列表: {list(sections)[:10]}{'...' if len(sections) > 10 else ''}")
        
    except Exception as e:
        print(f"❌ 检查失败: {e}")
        import traceback
        traceback.print_exc()
        
    finally:
        if initializer:
            initializer.close()

if __name__ == "__main__":
    main() 